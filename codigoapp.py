import streamlit as st
import re
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

# ==========================================
# 2. VARIABLES GLOBALES Y CONFIGURACIÓN
# ==========================================

# --- CONSTANTES DE ESTADOS PARA DROPDOWNS ---
CAAP_LOGICA_ESTADOS = ["No aplica","No iniciado, tiene CNCA vigente","No iniciado, tiene CNCA en curso","No iniciado, tiene CNCA vencida / No solicitada)","En curso","Vigente","Vencido"]
CAAF_LOGICA_ESTADOS = ["No aplica","No iniciado, tiene CAAP vigente","No iniciado, tiene CAAP en curso","No iniciado, tiene CAAP vencida / No solicitada)","En curso","Vigente","Vencido"]
ADA_DETALLE_ESTADOS = ["Seleccione...", "Prefactibilidad con Chi 0 para todos los permisos","Prefactibilidad superada", "Prefactibilidad vigente", "Prefactibilidad vencida", "No solicitada"]
RENPRE_DETALLE_ESTADOS = ["Seleccione...", "No aplica", "Esta inscripto y renueva", "Esta inscripto y no renovó", "Aplica pero no esta inscripto"]
FRECUENCIA_MUESTREO_OPTIONS = ["Seleccione...", "Anual", "Semestral", "Trimestral", "Mensual", "Continua", "N/A"]
AYSA_DETALLE_ESTADOS = ["Aplica","No aplica"]
RESIDUOS_ESPECIALES_STATUS_GENERAL = [
    "Seleccione...", "Empresa exenta", "No inscripta", "No Cumple con las DDJJ",
    "Cumple con las DDJJ"
]
CHE_DETALLE_ESTADOS = [
    "Seleccione...", "Obtuvo CHE alguna vez y esta en curso la renovacion", "No obtuvo CHE previamente y se encuentra en curso", "CHE vigente"
]
HIDRAULICA_DETALLE_ESTADOS = ["Seleccione...", "Tramite no iniciado porque tiene CHI 0", "No iniciado, Prefactibilidad vigente", "No iniciado, Prefactibilidad en curso", "No iniciado, Prefactibilidad no solicitada/vencida", "En curso","Vigente"]
VUELCO_DETALLE_ESTADOS = ["Seleccione...", "Tramite no iniciado porque tiene CHI 0", "No iniciado, Prefactibilidad vigente", "No iniciado, Prefactibilidad en curso", "No iniciado, Prefactibilidad no solicitada/vencida", "En curso","Vigente"]
EXPLOTACION_DETALLE_ESTADOS = ["Seleccione...", "Tramite no iniciado porque tiene CHI 0", "No iniciado, Prefactibilidad vigente", "No iniciado, Prefactibilidad en curso", "No iniciado, Prefactibilidad no solicitada/vencida", "En curso","Vigente"]
ACUMAR_DETALLE_ESTADOS = ["Seleccione...", "vigente NIA menor a 40", "vigente NIA mayor a 40", "no solicitada", "no aplica"]
SE_DETALLE_ESTADOS = ["Seleccione...", "No aplica", "Vigente", "En curso", "Vencida", "No solicitada"]
ULTIMO_CAA_ESTADOS = ["Tiene ultimo caa", "No tiene ultimo caa"]
ULTIMA_LEGA_ESTADOS = ["Tiene ultima lega", "No tiene ultima LEGA/PDEG", "Tiene ultimo permiso PDEG","La ultima lega es la vigente"]

# --- HALLAZGOS PREDEFINIDOS ---
HALLAZGOS_PREDEFINIDOS = {
    'Ambiental': [
        {
            'situacion': 'Durante el recorrido se evidenció que hubo un derrame de líquido con presencia de hidrocarburos sobre suelo absorbente.',
            'autoridad': 'Autoridad del Agua (ADA) o Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción ante una inspección de ADA o bien, MAPBA por el derrame, en relación con el Art. 103 de la Ley 12.257 - Código de Aguas de la Provincia de Buenos Aires, o el incumplimiento de la Resolución 3722/16 que establece que se debe informar cualquier eventualidad en sus operaciones que pueda impactar en el ambiente o generar preocupación en la comunidad.',
            'recomendacion': 'Identificar la posible causa para tomar acciones con el fin de evitar este tipo de derrames o salpicaduras sobre suelo absorbente. En caso de ocurrir, se debe dar aviso ante las autoridades pertinentes.'
        },
        {
            'situacion': 'Se constató que la empresa realiza almacenamiento de combustible en planta mediante sistemas aéreos, correspondientes al tanque de la sala de calderas y al tanque de bombas de la red de incendio. Sin embargo, dichos sistemas no se encuentran inscriptos.',
            'autoridad': 'Secretaría de Energía (SE).',
            'riesgo': 'Retrasar la emisión de permisos/habilitaciones si la autoridad detecta el tanque sin adecuar.',
            'recomendacion': 'Declarar los sistemas de almacenamiento ante la SE, inscribiéndolos en el Registro de Bocas de Expendio de Combustibles Líquidos (Res. 1102/04), para incorporarlos en las auditorías (Res. 404/94).'
        },
        {
            'situacion': 'Se observó un contenedor (bin) de 1000 Lts con combustible almacenado transitoriamente, dispuesto inadecuadamente a la intemperie, sobre suelo absorbente y sin identificación.',
            'autoridad': 'Secretaría de Energía (SE).',
            'riesgo': 'Retrasar la emisión de permisos/habilitaciones si la autoridad detecta el tanque sin adecuar.',
            'recomendacion': 'Retirar de planta o, si se planea mantener, adecuar el sistema para luego proceder con su habilitación ante la SE. Si se usa un batán, revisar la normativa particular de transporte.'
        },
        {
            'situacion': 'El cuarto de lavado de piezas se observó sucio y colapsado, con la rejilla de la cámara de contención desbordada, lo que provocó el estancamiento de líquido contaminado y charcos.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción durante una inspección de MAPBA.',
            'recomendacion': 'Realizar acciones para evitar este tipo de desbordes y verificar que la capacidad de almacenamiento de la cámara sea la adecuada para los volúmenes generados.'
        },
        {
            'situacion': 'Respecto al transformador eléctrico de vía húmeda, el análisis realizado fue de manera no oficial.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción ante una inspección de la MAPBA.',
            'recomendacion': 'Se recomienda realizar un análisis de manera oficial con protocolo de informe y certificado de cadena de custodia oficial.'
        },
        {
            'situacion': 'El sector donde se almacenan los contenedores para residuos no especiales (para transporte y reciclaje) no cuenta con protección contra las inclemencias climáticas.',
            'autoridad': 'Ministerio de Ambiente (MDA).',
            'riesgo': 'Infracción ante una inspección de la MDA.',
            'recomendacion': 'Se recomienda colocar los volquetes en un sector con protección contra las lluvias o colocar volquetes con tapas.'
        },
        {
            'situacion': 'Se tomó vista de una serie de bines y tambores sin identificar.',
            'autoridad': 'Ministerio de Ambiente (MDA).',
            'riesgo': 'Infracción ante una inspección de la MDA.',
            'recomendacion': 'Se recomienda identificar los bines y tambores observados definiendo si son residuos especiales (almacenar en depósito transitorio con etiquetas), materia prima (almacenar en depósito destinado para tal fin) o para devolución (definir y sectorizar un lugar).'
        },
        {
            'situacion': 'Los residuos almacenados dentro del depósito de residuos especiales no contaban con etiquetas identificatorias.',
            'autoridad': 'Ministerio de Ambiente (MDA).',
            'riesgo': 'Infracción ante una inspección de la MDA.',
            'recomendacion': 'Se recomienda incorporar etiquetas que contengan fecha de ingreso, categoría (Y) y peligrosidad (H) en todos los residuos almacenados dentro del depósito.'
        },
        {
            'situacion': 'El establecimiento no posee separación de los efluentes líquidos industriales y del proceso de refrigeración que permita evaluar la calidad previa a la CAyTM final, tal como lo solicita ACUMAR. Además, la CAyTM no cuenta con la placa para la clausura de vuelco.',
            'autoridad': 'Autoridad de Cuenca Matanza Riachuelo (ACUMAR) y Autoridad del Agua (ADA).',
            'riesgo': 'Infracción ante una inspección de la ACUMAR y ADA.',
            'recomendacion': 'Se recomienda evaluar la posibilidad de realizar 2 CAyTM (una para efluentes industriales y otra para efluentes de refrigeración) o enviar los efluentes de refrigeración a la PTEL e incorporar la placa para la clausura de vuelco.'
        },
        {
            'situacion': 'En la PTEL se observó una manguera utilizada para incorporar agua, práctica considerada dilución del efluente y que está prohibida.',
            'autoridad': 'Autoridad del Agua (ADA).',
            'riesgo': 'Infracción por parte de la Autoridad del Agua.',
            'recomendacion': 'Quitar las mangueras que se utilicen para verter agua dentro de la PTEL.'
        },
        {
            'situacion': 'El depósito de químicos de caldera (con contención y techo parciales) tenía envases almacenados sobre sectores donde el techo y la contención de derrames no cubrían.',
            'autoridad': 'Ministerio de Ambiente.',
            'riesgo': 'Infracción ante una inspección del Ministerio de Ambiente.',
            'recomendacion': 'Asegurar el almacenamiento en el sector adecuado del depósito o extender el techo y la contención para cubrir toda la superficie de la planta.'
        },
        {
            'situacion': 'Las adecuaciones implementadas en el Orificio Toma Muestra (OTM) para la LEGA no cumplen con los requisitos técnicos de las Res. 559/19 y Dec. 1074/18, ya que el diámetro de la instalación está por debajo del mínimo exigido.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción durante una inspección del organismo de control.',
            'recomendacion': 'Se recomienda proceder con la modificación del OTM a fin de garantizar el cumplimiento de las dimensiones mínimas estipuladas por la normativa.'
        },
        {
            'situacion': 'Se constató la presencia de un nuevo pozo de explotación hídrica no declarado formalmente ante la Autoridad del Agua (ADA). Esta captación no figura en los permisos de uso del recurso hídrico.',
            'autoridad': 'Autoridad del Agua (ADA).',
            'riesgo': 'Infracción por existencia de instalaciones no declaradas o por falta de condiciones del pozo.',
            'recomendacion': 'Declarar el nuevo pozo de explotación ante la ADA e incorporar un caudalímetro homologado para cumplir con los requerimientos de medición.'
        },
        {
            'situacion': 'Se constató la correcta señalización de equipos identificados como "libres de PCBs".',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Ante auditorías, es obligatorio que haya al menos un análisis de PCBs de los transformadores.',
            'recomendacion': 'Si no se tiene un monitoreo, realizarlo. Si el monitoreo tiene una fecha mayor a 3 años, realizarlo nuevamente para conocer el estatus actual.'
        },
        {
            'situacion': 'Presencia de baldes conteniendo residuos con materia orgánica, sin sistema de contención secundaria y sin identificación alguna.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Llamado de atención de las autoridades, o derrames descontrolados que terminen en el exterior de la planta.',
            'recomendacion': 'Se recomienda realizar la adecuación para contención de derrames, protección contra inclemencias climáticas y piso impermeable sin conexión con el sistema de pluviales.'
        },
        {
            'situacion': 'Tanque contenedor de ácido conectado a una manguera sin medidas de seguridad. El líquido era liberado directamente al suelo, sin contención secundaria ni sistemas de control de derrames.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA) y la Autoridad del Agua (ADA).',
            'riesgo': 'Infracción por eventualidades no declaradas (MAPBA) o por vertidos no declarados si se derivan al pluvial (ADA).',
            'recomendacion': 'Implementación de sistemas de contención de derrames en los puntos de carga y descarga, y un mejor guardado de la manguera.'
        },
        
    ],
    'Aparatos sometidos a presion ASP': [
         {
            'situacion': 'La planta no posee la cantidad mínima de foguistas habilitados.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción del MAPBA por incumplimiento al Art. 18 de la Res. 231/96 modificado por Art. 5 de la Res. 1126/07.',
            'recomendacion': 'Realizar la capacitación y habilitación correspondiente a operarios para cumplir con la cantidad mínima de foguistas de acuerdo a la cantidad de turnos.'
        },
        {
            'situacion': 'En la sala de calderas no se evidenció la presencia de protecciones y alarmas de detección automática de fuga de combustibles gaseosos y detectores de monóxido.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción del MAPBA por incumplimiento al Art. 18 de la Res. 231/96 modificado por Art. 5 de la Res. 1126/07.',
            'recomendacion': 'Se recomienda realizar la instalación de los elementos de seguridad previamente mencionados.'
        },
        {
            'situacion': 'En la sala de calderas no se evidenció la presencia del libro de seguimiento foliado de generadores de vapor, acorde al Apéndice 3 de la Resolución 1126/07, en el que se asienten todos los controles realizados, reparaciones solicitadas y/o realizadas, y todas las anormalidades detectadas con indicación de la fecha respectiva.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción del MAPBA por incumplimiento de la Resolución 1126/07.',
            'recomendacion': 'Se recomienda confeccionar el libro rubricado y colocarlo en la sala de calderas.'
        },
          {
            'situacion': 'La sala de calderas no contaba con detector de gas y monóxido de carbono.',
            'autoridad': 'Ministerio de Ambiente.',
            'riesgo': 'Infracción ante una inspección del Ministerio de Ambiente.',
            'recomendacion': 'Se recomienda avanzar en la colocación del detector.'
        },
          {
            'situacion': 'Durante el relevamiento se observó que, respecto de los ASP, no se encontraba identificado el destino final de las purgas asociadas.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción durante una inspección del organismo de control.',
            'recomendacion': 'Se sugiere la verificación de la salida de las mismas, considerando que podrían contener mezclas de aceite y agua, residuo clasificado como especial según la normativa vigente.'
        },
    ],
    'Residuos Especiales': [
       {
            'situacion': 'Se observan residuos especiales acopiados fuera del depósito sin cobertura ante inclemencias climáticas.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción al Art. 3 inciso A de la Resolución 592/00.',
            'recomendacion': 'Almacenar los residuos especiales en el depósito transitorio con la cobertura y contención adecuadas, conforme a la Resolución 592/00.'
        },
          {
            'situacion': 'No fue posible acceder al depósito de residuos especiales debido a que se encontraba en proceso de reubicación, impidiendo verificar el cumplimiento de las disposiciones técnicas.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción por no poseer depósito de residuos especiales conforme a la Resolución 592/00.',
            'recomendacion': 'Debe realizarse de manera urgente la adecuación de un sector para el almacenamiento de residuos especiales conforme a la Resolución 592/00.'
        },
       {
            'situacion': 'Algunos residuos almacenados en el depósito de residuos especiales no se encuentran separados por pasillos de 1 metro, impidiendo la visualización de los residuos posteriores.',
            'autoridad': 'Ministerio de Ambiente de la Provincia de Buenos Aires (MAPBA).',
            'riesgo': 'Infracción durante una inspección de MAPBA por incumplimiento de la Resolución 592/00.',
            'recomendacion': 'Ordenar los residuos para facilitar la verificación y contabilización ante una inspección. Adicionalmente, revisar que se cumpla con el etiquetado de la totalidad de residuos.'
        },
        
    ],
}

# --- MARCADORES_CONDICIONALES  ---
MARCADORES_CONDICIONALES = {
    "HABILITACION_MUNICIPAL": {
        "cumple": {"start": "{INICIO_HAB_CUMPLE}", "end": "{FIN_HAB_CUMPLE}"}, "no cumple": {"start": "{INICIO_HAB_NO_CUMPLE}", "end": "{FIN_HAB_NO_CUMPLE}"},
        "parcial": {"start": "{INICIO_HAB_PARCIAL}", "end": "{FIN_HAB_PARCIAL}"}
    },
    "CNCA_STATUS": {
        "vigente": {"start": "{INICIO_CNCA_VIGENTE}", "end": "{FIN_CNCA_VIGENTE}"}, "superada (esta en curso el CAA)": {"start": "{INICIO_CNCA_SUPERADA}", "end": "{FIN_CNCA_SUPERADA}"},
        "vencida": {"start": "{INICIO_CNCA_VENCIDA}", "end": "{FIN_CNCA_VENCIDA}"}, "no solicitada": {"start": "{INICIO_CNCA_NO_SOLICITADA}", "end": "{FIN_CNCA_NO_SOLICITADA}"},
        "en curso": {"start": "{INICIO_CNCA_EN_CURSO}", "end": "{FIN_CNCA_EN_CURSO}"}
    },
    "CAAP_STATUS": {
        "No iniciado, tiene CNCA vigente": {"start": "{INICIO_CAAP_NO_INICIADO_CNCA_VIGENTE}", "end": "{FIN_CAAP_NO_INICIADO_CNCA_VIGENTE}"},
        "No iniciado, tiene CNCA en curso": {"start": "{INICIO_CAAP_NO_INICIADO_CNCA_ENCURSO}", "end": "{FIN_CAAP_NO_INICIADO_CNCA_ENCURSO}"},
        "No iniciado, tiene CNCA vencida / No solicitada)": {"start": "{INICIO_CAAP_NO_INICIADO_CNCA_VENCIDA_NO_SOLICITADA}", "end": "{FIN_CAAP_NO_INICIADO_CNCA_VENCIDA_NO_SOLICITADA}"},
        "En curso": {"start": "{INICIO_CAAP_EN_CURSO}", "end": "{FIN_CAAP_EN_CURSO}"},
        "Vigente": {"start": "{INICIO_CAAP_VIGENTE}", "end": "{FIN_CAAP_VIGENTE}"},
        "Vencido": {"start": "{INICIO_CAAP_VENCIDO}", "end": "{FIN_CAAP_VENCIDO}"},
        "No aplica": {"start": None,  "end": None}
    },
    "CAAF_STATUS": {
        "En curso": {"start": "{INICIO_CAAF_EN_CURSO}", "end": "{FIN_CAAF_EN_CURSO}"}, "vigente": {"start": "{INICIO_CAAF_VIGENTE}", "end": "{FIN_CAAF_VIGENTE}"},
        "No iniciado, CAAP en curso": {"start": "{INICIO_CAAF_NO_INICIADO_CAAP_EN_CURSO}", "end": "{FIN_CAAF_NO_INICIADO_CAAP_EN_CURSO}"},
        "No iniciado, CAAP vencido": {"start": "{INICIO_CAAF_NO_INICIADO_CAAP_VENCIDO}", "end": "{FIN_CAAF_NO_INICIADO_CAAP_VENCIDO}"},
        "Vencido": {"start": "{INICIO_CAAF_VENCIDO}", "end": "{FIN_CAAF_VENCIDO}"},
        "Eliminar todo": {"start": "{ELIMINAR_TODO_CAAF}", "end": "{FIN_ELIMINAR_TODO_CAAF}"} # Special case for "No aplica"
    },
    "ULTIMO_CAA_STATUS": {
        "Tiene ultimo caa": {"start": "{INICIO_ULTIMO_CAA_OBTENIDO}", "end": "{FIN_ULTIMO_CAA_OBTENIDO}"},
        "No tiene ultimo caa": {"start": "{INICIO_CAA_NUNCA_OBTENIDO}", "end": "{FIN_CAA_NUNCA_OBTENIDO}"}
    },
    "RENOVACION_CAA_STATUS": {
        "En curso": {"start": "{INICIO_RENOVACION_CAA_ENCURSO}", "end": "{FIN_RENOVACION_CAA_ENCURSO}"},
        "Finalizada": {"start": "{INICIO_RENOVACION_CAA_FINALIZADA}", "end": "{FIN_RENOVACION_CAA_FINALIZADA}"},
        "No iniciada": {"start": "{INICIO_RENOVACION_NO_INICIADA}", "end": "{FIN_RENOVACION_NO_INICIADA}"},
         "No aplica":  { "start": None, "end": None}
    },
    "LEGA_STATUS": {
        "vigente": {"start": "{INICIO_LEGA_VIGENTE}", "end": "{FIN_LEGA_VIGENTE}"}, "en_curso": {"start": "{INICIO_LEGA_EN_CURSO}", "end": "{FIN_LEGA_EN_CURSO}"},
        "vencida": {"start": "{INICIO_LEGA_VENCIDA}", "end": "{FIN_LEGA_VENCIDA}"}
    },
    "RESIDUOS_ESPECIALES_STATUS": {
        "Empresa exenta": {"start": "{INICIO_EMPRESA_EXCENTA}", "end": "{FIN_EMPRESA_EXCENTA}"},
        "No inscripta": {"start": "{INICIO_RREE_NO_INSCRIPTA}", "end": "{FIN_RREE_NO_INSCRIPTA}"},
        "No Cumple con las DDJJ": {"start": "{INICIO_RREE_NOCUMPLE_DDJJ}", "end": "{FIN_RREE_NOCUMPLE_DDJJ}"},
        "Cumple con las DDJJ": {"start": "{INICIO_RREE_CUMPLE_DDJJ}", "end": "{FIN_RREE_CUMPLE_DDJJ}"}
    },
    "CHE_STATUS": {
        "Obtuvo CHE alguna vez y esta en curso la renovacion": {"start": "{INICIO_RREE_OBTUVO_CHE}", "end": "{FIN_RREE_OBTUVO_CHE}"},
        "Vigente": {"start": "{INICIO_RREE_CHE_VIGENTE}", "end": "{FIN_RREE_CHE_VIGENTE}"},
        "No obtuvo CHE previamente y se encuentra en curso": {"start": "{INICIO_RREE_CHE_EN_CURSO}", "end": "{FIN_RREE_CHE_EN_CURSO}"}
    },
    "GIRSU_STATUS": {
        "Aplica": {"start": "{INICIO_APLICA_GIRSU}", "end": "{FIN_APLICA_GIRSU}"},
        "No aplica": {"start": "{INICIO_NOAPLICA_GIRSU}", "end": "{FIN_NOAPLICA_GIRSU}"},
        "Aplica y realizo la presentacion": {"start": "{INICIO_APLICA_GIRSU_CUMPLE}", "end": "{FIN_APLICA_GIRSU_CUMPLE}"},
        "Aplica y no realizo la presentacion": {"start": "{INICIO_APLICA_GIRSU_NOCUMPLE}", "end": "{FIN_APLICA_GIRSU_NOCUMPLE}"}
    },
    "PATOGENICOS_STATUS":{
        "Inscripto": {"start": "{INICIO_INSCRIPTA_PATOGENICOS}", "end": "{FIN_INSCRIPTA_PATOGENICOS}"},
        "No inscripto": {"start": "{INICIO_NO_INSCRIPTA_PATOGENICOS}", "end": "{FIN_NO_INSCRIPTA_PATOGENICOS}"},
        "No aplica": {"start": "{INICIO_NO_APLICA_PATOGENICOS}", "end": "{FIN_NO_APLICA_PATOGENICOS}"}
    },
    "ASP_STATUS": {
        "Finalizada": {"start": "{INICIO_PRESENTACION_ASP_FINALIZADA}", "end": "{FIN_PRESENTACION_ASP_FINALIZADA}"},
        "Caratulada": {"start": "{INICIO_PRESENTACION_ASP_CARATULADA}", "end": "{FIN_PRESENTACION_ASP_CARATULADA}"},
        "No Presentado": {"start": "{INICIO_ASP_NO_PRESENTADO}", "end": "{FIN_ASP_NO_PRESENTADO}"}
    },
    "VALVULAS_CALIBRACION_STATUS": {
        "Cumple": {"start": "{INICIO_CALIBRACION_ASP_CUMPLE}", "end": "{FIN_CALIBRACION_ASP_CUMPLE}"},
        "No Cumple": {"start": "{INICIO_CALIBRACION_ASP_NOCUMPLE}", "end": "{FIN_CALIBRACION_ASP_NOCUMPLE}"}
    },
    "ADA_STATUS": {
         "Prefactibilidad con Chi 0 para todos los permisos": {"start": "{INICIO_PREFA_TODOS_CHI0}", "end": "{FIN_PREFA_TODOS_CHI0}"},
         "Prefactibilidad vigente": {"start": "{INICIO_PREFA_VIGENTE}", "end": "{FIN_PREFA_VIGENTE}"},
         "Prefactibilidad superada": {"start": "{INICIO_PREFA_SUPERADA}", "end": "{FIN_PREFA_SUPERADA}"},
         "Prefactibilidad vencida": {"start": "{INICIO_PREFA_VENCIDA}", "end": "{FIN_PREFA_VENCIDA}"},
         "No solicitada": {"start": "{INICIO_PREFA_NO_SOLICITADA}", "end": "{FIN_PREFA_NO_SOLICITADA}"}
    },
    "RENPRE_STATUS": {
        "No aplica": {"start": "{INICIO_NO_APLICA}", "end": "{FIN_NO_APLICA}"},
        "Esta inscripto y renueva": {"start": "{INICIO_APLICA_INSCRIPTO_RENUEVA}", "end": "{FIN_APLICA_INSCRIPTO_RENUEVA}"},
        "Esta inscripto y no renovó": {"start": "{INICIO_APLICA_INSCRIPTO_NO_RENOVO}", "end": "{FIN_APLICA_INSCRIPTO_NO_RENOVO}"},
        "Aplica pero no esta inscripto": {"start": "{INICIO_APLICA_NO_INSCRIPTO}", "end": "{FIN_APLICA_NO_INSCRIPTO}"}
    },
    "PLANDEADECUACIONACUMAR_STATUS": {
        "Plan de adecuacion en curso": {"start": "{INICIO_ACUMAR_PLAN_ADECUACION_EN_CURSO}", "end": "{FIN_ACUMAR_PLAN_ADECUACION_EN_CURSO}"},
        "No aplica Plan de adecuacion": {"start": "{INICIO_ACUMAR_PLAN_ADECUACION_NO_APLICA}", "end": "{FIN_ACUMAR_PLAN_ADECUACION_NO_APLICA}"}
    },
    "SEGURO_STATUS": {
        "Vigente": {"start": "{INICIO_POLIZA_VIGENTE}", "end": "{FIN_POLIZA_VIGENTE}"}, "Vencida": {"start": "{INICIO_POLIZA_VENCIDA}", "end": "{FIN_POLIZA_VENCIDA}"},
        "Nunca tuvo pero le aplica": {"start": "{INICIO_NUNCA_TUVO_POLIZA_APLICA}", "end": "{FIN_NUNCA_TUVO_POLIZA_APLICA}"},
        "Nunca tuvo y no aplica": {"start": "{INICIO_NUNCA_TUVO_POLIZA_VERIFICAR}", "end": "{FIN_NUNCA_TUVO_POLIZA_VERIFICAR}"}
    },
    "ACUMAR_STATUS": {
        "vigente NIA menor a 40": {"start": "{INICIO_ACUMAR_VIGENTE_NIA_MENOR40}", "end": "{FIN_ACUMAR_VIGENTE_NIA_MENOR40}"},
        "vigente NIA mayor a 40": {"start": "{INICIO_ACUMAR_VIGENTE_NIA_MAYOR40}", "end": "{FIN_ACUMAR_VIGENTE_NIA_MAYOR40}"},
        "no aplica": {"start": "{INICIO_ACUMAR_NO_APLICA}", "end": "{FIN_ACUMAR_NO_APLICA}"},
        "no solicitada": {"start": "{INICIO_ACUMAR_NO_SOLICITADA}", "end": "{FIN_ACUMAR_NO_SOLICITADA}"}
    },
    "INSCRIPCION_1102": {
        "Inscripto": {"start": "{INICIO_INSCRIPTA_1102}", "end": "{FIN_INSCRIPTA_1102}"},
        "No inscripto": {"start": "{INICIO_NOINSCRIPTA_1102}", "end": "{FIN_NOINSCRIPTA_1102}"},
        "No aplica": {"start": "{INICIO_NOAPLICA_1102}", "end": "{FIN_NOAPLICA_1102}"}

    },
    "AUDITORIA_404": {
        "Realizo": {"start": "{INICIO_REALIZO_AUDITORIA}", "end": "{FIN_REALIZO_AUDITORIA}"},
        "No realizo": {"start": "{INICIO_NO_REALIZO_AUDITORIA}", "end": "{FIN_NO_REALIZO_AUDITORIA}"},
        "No aplica": {"start": "{INICIO_NOAPLICA_AUDITORIA}", "end": "{FIN_NOAPLICA_AUDITORIA}"},
      "No inscripto, no realiza": {"start": "{INICIO_NO_INSCRIPTA_NO_AUDITORIA}", "end": "{FIN_NO_INSCRIPTA_NO_AUDITORIA}"}
    },
    "INSCRIPCION_277": {
        "Aplica": {"start": "{INICIO_APLICA_AUDITORIA277}", "end": "{FIN_APLICA_AUDITORIA277}"},
        "No aplica": {"start": "{INICIO_NOAPLICA_277}", "end": "{FIN_NOAPLICA_277}"}

    },
   
    "ULTIMA_LEGA_STATUS": {
        "Tiene ultima lega": {"start": "{INICIO_ULTIMA_LEGA_OBTENIDA}", "end": "{FIN_ULTIMA_LEGA_OBTENIDA}"},
        "No tiene ultima LEGA/PDEG": {"start": "{INICIO_LEGA_NUNCA_OBTENIDA}", "end": "{FIN_LEGA_NUNCA_OBTENIDA}"},
        "Tiene ultimo permiso PDEG": {"start": "{INICIO_ULTIMO_PDEG_OBTENIDO}", "end": "{FIN_ULTIMO_PDEG_OBTENIDO}"},
        "La ultima lega es la vigente":{"start":None, "end": None}
        
    },
     "HIDRAULICA_STATUS": {
        "Tramite no iniciado porque tiene CHI 0": {"start": "{INICIO_CONSTANCIA_HIDRAULICA_NO_INICIADA_POR_CHI0}", "end": "{FIN_CONSTANCIA_HIDRAULICA_NO_INICIADA_POR_CHI0}"},
        "No iniciada, tiene Prefactibilida vigente": {"start": "{INICIO_CONSTANCIA_HIDRAULICA_NO_INICIADA_PREFA_VIGENTE}", "end": "{FIN_CONSTANCIA_HIDRAULICA_NO_INICIADA_PREFA_VIGENTE}"},
        "No iniciada, Prefactibilidad en curso": {"start": "{INICIO_CONSTANCIA_HIDRAULICA_NO_INICIADA_PREFA_EN_CURSO}", "end": "{FIN_CONSTANCIA_HIDRAULICA_NO_INICIADA_PREFA_EN_CURSO}"},
        "No iniciada, Prefactibilidad vencida o no iniciada": {"start": "{INICIO_CONSTANCIA_HIDRAULICA_NO_INICIADA_PREFA_VENCIDA_NO_INICIADA}", "end": "{FIN_CONSTANCIA_HIDRAULICA_NO_INICIADA_PREFA_VENCIDA_NO_INICIADA}"},
        "En curso": {"start": "{INICIO_CONSTANCIA_HIDRAULICA_EN_CURSO}", "end": "{FIN_CONSTANCIA_HIDRAULICA_EN_CURSO}"},
        "Vigente": {"start": "{INICIO_CONSTANCIA_HIDRAULICA_VIGENTE}", "end": "{FIN_CONSTANCIA_HIDRAULICA_VIGENTE}"}
    },
    "ULTIMOHIDRAULICA_STATUS": {
        "tiene ultimo permiso hidraulica": {"start": "{INICIO_ULTIMA_CONSTANCIA_HIDRAULICA_OBTENIDO}", "end": "{FIN_ULTIMA_CONSTANCIA_HIDRAULICA_OBTENIDO}"},
        "no tiene ultimo permiso de hidraulica": {"start": None,  "end": None}
    },
    "VUELCO_STATUS": {
        "Tramite no iniciado porque tiene CHI 0": {"start": "{INICIO_PERMISO_VUELCO_NO_INICIADO_POR_CHI0}", "end": "{FIN_PERMISO_VUELCO_NO_INICIADO_POR_CHI0}"},
        "No iniciada, tiene Prefactibilida vigente": {"start": "{INICIO_PERMISO_VUELCO_NO_INICIADO_PREFA_VIGENTE}", "end": "{FIN_PERMISO_VUELCO_NO_INICIADO_PREFA_VIGENTE}"},
        "No iniciada, Prefactibilidad en curso": {"start": "{INICIO_PERMISO_VUELCO_NO_INICIADO_PREFA_EN_CURSO}", "end": "{FIN_PERMISO_VUELCO_NO_INICIADO_PREFA_EN_CURSO}"},
        "No iniciada, Prefactibilidad vencida o no iniciada": {"start": "{INICIO_PERMISO_VUELCO_NO_INICIADO_PREFA_VENCIDA_NO_INICIADA}", "end": "{FIN_PERMISO_VUELCO_NO_INICIADO_PREFA_VENCIDA_NO_INICIADA}"},
        "En curso": {"start": "{INICIO_PERMISO_VUELCO_EN_CURSO}", "end": "{FIN_PERMISO_VUELCO_EN_CURSO}"},
        "Vigente": {"start": "{INICIO_PERMISO_VUELCO_VIGENTE}", "end": "{FIN_PERMISO_VUELCO_VIGENTE}"}
    },
    "ULTIMOVUELCO_STATUS": {
        "tiene ultimo permiso vuelco": {"start": "{INICIO_ULTIMO_PERMISO_VUELCO_OBTENIDO}", "end": "{FIN_ULTIMO_PERMISO_VUELCO_OBTENIDO}"},
        "no tiene ultimo permiso de vuelco": { "start": None, "end": None}
    },
    "EXPLOTACION_STATUS": {
        "Tramite no iniciado porque tiene CHI 0": {"start": "{INICIO_PERMISO_EXPLOTACION_NO_INICIADO_POR_CHI0}", "end": "{FIN_PERMISO_ EXPLOTACION_NO_INICIADO_POR_CHI0}"},
        "No iniciada, tiene Prefactibilida vigente": {"start": "{INICIO_PERMISO_EXPLOTACION_NO_INICIADO_PREFA_VIGENTE}", "end": "{FIN_PERMISO_EXPLOTACION_NO_INICIADO_PREFA_VIGENTE}"},
        "No iniciada, Prefactibilidad en curso": {"start": "{INICIO_PERMISO_EXPLOTACION_NO_INICIADO_PREFA_EN_CURSO}", "end": "{FIN_PERMISO_EXPLOTACION_NO_INICIADO_PREFA_EN_CURSO}"},
        "No iniciada, Prefactibilidad vencida o no iniciada": {"start": "{INICIO_PERMISO_EXPLOTACION_NO_INICIADO_PREFA_VENCIDA_NO_INICIADA}", "end": "{FIN_PERMISO_EXPLOTACION_NO_INICIADO_PREFA_VENCIDA_NO_INICIADA}"},
        "En curso": {"start": "{INICIO_PERMISO_EXPLOTACION_EN_CURSO}", "end": "{FIN_PERMISO_EXPLOTACION_EN_CURSO}"},
        "Vigente": {"start": "{INICIO_PERMISO_EXPLOTACION_VIGENTE}", "end": "{FIN_PERMISO_EXPLOTACION_VIGENTE}"}
    },
    "ULTIMOEXPLOTACION_STATUS": {
        "tiene ultimo permiso explotacion": {"start": "{INICIO_ULTIMO_PERMISO_EXPLOTACION_OBTENIDO}", "end": "{FIN_ULTIMO_PERMISO_EXPLOTACION_OBTENIDO}"},
        "no tiene ultimo permiso de explotacion":  { "start": None, "end": None}
    },
    "AYSA_STATUS": {
        "Aplica": {"start": "{INICIO_APLICA_AYSA}", "end": "{FIN_APLICA_AYSA}"},
        "No aplica": {"start": "{INICIO_AYSA_NO_APLICA}", "end": "{FIN_AYSA_NO_APLICA}"},
}
}
# --- INICIALIZACIÓN DEL ESTADO DE SESIÓN ---
if 'hallazgos_widgets_list' not in st.session_state:
    st.session_state['hallazgos_widgets_list'] = []

# PARA EL PLAN DE MONITOREO:
if 'muestreo_filas_datos' not in st.session_state:
   
    st.session_state['muestreo_filas_datos'] = [{
        'recurso': '', 
        'organismo': '', 
        'puntos': '', 
        'parametros': '', 
        'frecuencia': FRECUENCIA_MUESTREO_OPTIONS[0]
    }]

if 'datos_generales' not in st.session_state:
    st.session_state['datos_generales'] = {}
# También inicializamos otras variables si no existen
if 'datos_generales' not in st.session_state:
    st.session_state['datos_generales'] = {}
# ==========================================
# 3. FUNCIONES DE PROCESAMIENTO DOCX
# ==========================================

def agregar_hallazgo_formateado_al_doc(doc, index, situacion, autoridad, riesgo, recomendacion):
    if not situacion.strip(): return

    p_title = doc.add_paragraph()
    run_title = p_title.add_run(f"Observación de campo # {index}")
    run_title.bold = True
    run_title.underline = True
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    fields = [
      ("Situación:", situacion), ("Autoridad:", autoridad),
        ("Riesgo:", riesgo), ("Recomendación:", recomendacion)
    ]

    doc.add_paragraph() # Separator

    for label, value in fields:
        if value and value.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_label = p.add_run(label + " ")
            run_label.bold = True
            p.add_run(value)

def find_paragraphs_to_remove(doc, selected_state, situation_type):
    paragraphs_to_remove = []
    markers_config = MARCADORES_CONDICIONALES.get(situation_type, {})

    # El marcador que el usuario eligió (el que NO se debe borrar)
    selected_start = markers_config.get(selected_state, {}).get('start')

    # Todos los marcadores posibles para esta situación
    all_starts = {cfg.get('start') for cfg in markers_config.values() if cfg.get('start')}
    all_ends = {cfg.get('end') for cfg in markers_config.values() if cfg.get('end')}

    in_unselected_section = False

    for p in doc.paragraphs:
        text = p.text.strip()

        # 1. Encontró un inicio
        if text in all_starts:
            paragraphs_to_remove.append(p) # El marcador siempre se borra
            
            if text == selected_start:
                # ¡Es el que quiero! NO borro lo que viene a continuación
                in_unselected_section = False
            else:
                # No es el mío, activo el borrado
                in_unselected_section = True
            continue

        # 2. Encontró un fin
        if text in all_ends:
            paragraphs_to_remove.append(p) # El marcador siempre se borra
            in_unselected_section = False # Se apaga el borrado (lo que sigue es FIJO)
            continue

        # 3. Procesar párrafos intermedios
        if in_unselected_section:
            paragraphs_to_remove.append(p)
        
        # Si in_unselected_section es False, el párrafo se mantiene intacto (Texto Fijo)

    return paragraphs_to_remove

def reemplazar_marcadores(doc, user_data):
    def process_container_for_replacements_and_highlights(container):
        paragraphs_to_iterate = []
        if hasattr(container, 'paragraphs'): # It's a table cell
            paragraphs_to_iterate = container.paragraphs
        else: # It's a paragraph
            paragraphs_to_iterate = [container]

        for p in paragraphs_to_iterate:
            if not p.runs:
                continue

            first_run_format = {}
            if p.runs:
                run = p.runs[0]
                first_run_format = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color_rgb': run.font.color.rgb if run.font.color else None
                }

            combined_text = "".join([run.text for run in p.runs])
            modified_text = combined_text
            replaced_in_paragraph = False

            for marker, value in user_data.items():
                placeholder_pattern = re.compile(r'{\s*' + re.escape(marker) + r'\s*}')

                if placeholder_pattern.search(modified_text):
                    if str(value).strip() in ['N/A', '0', '']:
                        modified_text = placeholder_pattern.sub('', modified_text)
                        replaced_in_paragraph = True
                    elif str(value).strip():
                        modified_text = placeholder_pattern.sub(str(value), modified_text)
                        replaced_in_paragraph = True

            if replaced_in_paragraph:
                for i in range(len(p.runs) -1, -1, -1):
                    p.runs[i]._element.getparent().remove(p.runs[i]._element)
                new_run = p.add_run(modified_text)
                if first_run_format:
                    new_run.bold = first_run_format['bold']
                    new_run.italic = first_run_format['italic']
                    new_run.underline = first_run_format['underline']
                    if first_run_format['font_name']: new_run.font.name = first_run_format['font_name']
                    if first_run_format['font_size']: new_run.font.size = first_run_format['font_size']
                    if first_run_format['font_color_rgb']: new_run.font.color.rgb = first_run_format['font_color_rgb']

            for run in p.runs:
                if re.search(r'{\s*[A-Z_]+\s*}', run.text):
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for p in doc.paragraphs:
        process_container_for_replacements_and_highlights(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_container_for_replacements_and_highlights(cell)

def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def insertar_tabla_manual_dinamica(doc, lista_de_filas_datos):
    target_text = "Plan de monitoreos"
    for p in doc.paragraphs:
        if target_text in p.text:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            headers = ['recurso', 'organismo', 'puntos', 'parametros', 'frecuencia']
            hdr_cells = table.rows[0].cells
            for i, name in enumerate(headers):
                hdr_cells[i].text = name
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for fila in lista_de_filas_datos:
                recurso = fila.get('recurso', '').strip()
                if recurso:
                    row = table.add_row().cells
                    row[0].text = recurso
                    row[1].text = fila.get('organismo', '')
                    row[2].text = fila.get('puntos', '')
                    row[3].text = fila.get('parametros', '')
                    frec = fila.get('frecuencia', 'Seleccione...')
                    row[4].text = frec if frec != 'Seleccione...' else ""

                    for cell in row:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(10)
            break

# ==========================================
# 4. STREAMLIT APP LAYOUT AND LOGIC
# ==========================================

st.set_page_config(layout="wide")
st.title("📝 Generador de Informes Ambientales")

# Initialize session state variables
if 'uploaded_file' not in st.session_state:
    st.session_state['uploaded_file'] = None

if 'hallazgos_widgets_list' not in st.session_state:
    # Initial empty finding dictionary structure
    st.session_state['hallazgos_widgets_list'] = [{
     'situacion': '', 'autoridad': '', 'riesgo': '', 'recomendacion': ''
    }]

if 'muestreo_filas_datos' not in st.session_state:
    # Initial empty monitoring row dictionary structure
    st.session_state['muestreo_filas_datos'] = [{
        'Recurso': '', 'Organismo': '', 'Puntos': '',
        'Parametros': '', 'Frecuencia': 'Seleccione...'
    }]

# File Uploader in sidebar
st.sidebar.header("Cargar Plantilla")
uploaded_file = st.sidebar.file_uploader("Sube tu plantilla DOCX", type=["docx"])
if uploaded_file is not None:
    st.session_state['uploaded_file'] = uploaded_file
    st.sidebar.success("✅ Plantilla cargada correctamente.")
elif st.session_state['uploaded_file'] is None:
    st.sidebar.warning("Por favor, sube una plantilla DOCX para empezar.")

# Tabs for navigation
tab_titles = [
    '1. Info General', '2. CNCA, CAAP, CAAF', '3. LEGA y Monitoreos Ambientales',
    '4. Autoridad del Agua', '5. Otros requisitos de Agua','6. Otros Organismos', '7. Residuos y ASP',
    '8. Hallazgos y Recomendaciones'
]
tabs = st.tabs(tab_titles)

# --- 1. Información General y Habilitación Municipal ---
with tabs[0]:
    st.header("📍 Información Básica")
    st.session_state['NOMBRE_EMPRESA'] = st.text_input("Razón Social:", value="EMPRESA S.A.", key="main_NOMBRE_EMPRESA")
    st.session_state['NOMBRE_PLANTA'] = st.text_input("Planta Industrial:", value="Planta Industrial", key="main_NOMBRE_PLANTA")
    st.session_state['MES_AUDITORIA'] = st.text_input("Mes de relevamiento (Mes/Año):", value="Marzo 2026", key="main_MES_AUDITORIA")
    st.session_state['DIRECCION_PLANTA'] = st.text_input("Dirección:", value="", key="main_DIRECCION")
    st.session_state['MUNICIPIO_EMPRESA'] = st.text_input("Municipio:", value="", key="main_MUNICIPIO")
    st.session_state['RUBRO_EMPRESA'] = st.text_input("Rubro:", value="", key="main_RUBRO")

    st.header("✅ Habilitación Municipal")
    st.session_state['HAB_STATUS'] = st.selectbox("Estado Habilitación:", options=['Seleccione...', 'cumple', 'no cumple', 'parcial'], key="hab_STATUS")

    hab_disabled = st.session_state['HAB_STATUS'] == 'Seleccione...'
    st.session_state['FECHA_HABILITACION'] = st.text_input("Fecha de obtención de habilitación:", "dd/mm/aaaa", disabled=hab_disabled, key="hab_FECHA")
    st.session_state['EXPEDIENTE_HABILITACION'] = st.text_input("Nº de expediente:", "N/A", disabled=hab_disabled, key="hab_EXPEDIENTE")
    st.session_state['OBSERVACION_HAB_MUNICIPAL'] = st.text_area("Observaciones extra habilitación municipal:", "", disabled=hab_disabled, key="hab_OBSERVACION")

# --- 2. CNCA, CAAP, CAAF y Historial ---
with tabs[1]:
    st.header("📄 CNCA (Clasificación Nivel)")
    st.session_state['CNCA_STATUS'] = st.selectbox("Estado CNCA:", options=["No aplica", "vigente", "superada (esta en curso el CAA)", "vencida", "no solicitada", "en curso"], key="cnca_STATUS")

    cnca_disabled = st.session_state['CNCA_STATUS'] == 'No aplica'
    st.session_state['FECHA_CNCA'] = st.text_input("Fecha de obtención de CNCA:", "dd/mm/aaaa", disabled=cnca_disabled, key="cnca_FECHA")
    st.session_state['VENCIMIENTO_CNCA'] = st.text_input("Vencimiento de CNCA:", "dd/mm/aaaa", disabled=cnca_disabled, key="cnca_VENCIMIENTO")
    st.session_state['EXPEDIENTE_CNCA'] = st.text_input("Expediente de la CNCA:", "N/A", disabled=cnca_disabled, key="cnca_EXPEDIENTE")
    st.session_state['CATEGORIA_CNCA'] = st.text_input("Categoría:", "primera/segunda/tercera", disabled=cnca_disabled, key="cnca_CATEGORIA")
    st.session_state['PUNTOS_CNCA'] = st.text_input("Puntos (solo numero):", "ej:25", disabled=cnca_disabled, key="cnca_PUNTOS")
    st.session_state['DISPO_CNCA'] = st.text_input("Disposición de CNCA:", "N/A", disabled=cnca_disabled, key="cnca_DISPO")
    st.session_state['OBSERVACIONES_CNCA'] = st.text_area("Observaciones extra CNCA:", "", disabled=cnca_disabled, key="cnca_OBSERVACIONES")

    st.header("🌳 CAAP (Aptitud Ambiental - Fase II)")
    st.selectbox("Situación CAAP (Fase II):", options=["No aplica","No iniciado, tiene CNCA vigente","No iniciado, tiene CNCA en curso","No iniciado, tiene CNCA vencida / No solicitada)","En curso","Vigente","Vencido"], key="caap_CAAP_LOGICA_ESTADOS")
    caap_disabled = st.session_state.get('caap_CAAP_LOGICA_ESTADOS', '') == 'No aplica'
    st.session_state['FECHA_CAAP'] = st.text_input("Fecha de obtención del CAAP:", "N/A", disabled=caap_disabled, key="caap_FECHA")
    st.session_state['EXP_CAAP'] = st.text_input("Expediente del CAAP:", "N/A", disabled=caap_disabled, key="caap_EXP")
    st.session_state['DISPO_CAAP'] = st.text_input("Disposición del CAAP:", "N/A", disabled=caap_disabled, key="caap_DISPO")
    st.session_state['VIGENCIA_CAAP'] = st.text_input("Plazo de vigencia (años):", "N/A", disabled=caap_disabled, key="caap_VIGENCIA")
    st.session_state['VTO_CAAP'] = st.text_input("Vencimiento del CAAP:", "N/A", disabled=caap_disabled, key="caap_VTO")
    st.session_state['ESTADO_PORTAL_CAAP'] = st.text_input("Estado del CAAP en el portal:", "N/A", disabled=caap_disabled, key="caap_ESTADO_PORTAL")
    st.session_state['OBSERVACIONES_EXTRA_CAAP'] = st.text_area("Observaciones extra CAAP:", "", disabled=caap_disabled, key="caap_OBSERVACIONES")

    st.header("🌳 CAAF (Aptitud Ambiental - Fase III)")
    st.session_state['CAAF_LOGICA_ESTADOS'] = st.selectbox("Situación CAAF (Fase III):", options=["No aplica","No iniciado, tiene CAAP vigente","No iniciado, tiene CAAP en curso","No iniciado, tiene CAAP vencida / No solicitada)","En curso","Vigente","Vencido"], key="caaf_CAAF_LOGICA_ESTADOS")
    
    st.session_state['FECHA_CAAF'] = st.text_input("Fecha de obtención del CAAF:", "N/A", key="caaf_FECHA")
    st.session_state['EXPEDIENTE_CAAF'] = st.text_input("Expediente del CAAF:", "N/A",  key="caaf_EXPEDIENTE")
    st.session_state['DISPO_CAAF'] = st.text_input("Disposición de CAAF:", "N/A",  key="caaf_DISPO")
    st.session_state['ESTADO_PORTAL_CAAF'] = st.text_input("Estado del CAAF en el portal:", "N/A",  key="caaf_ESTADO_PORTAL")
    st.session_state['VENCIMIENTO_CAAF'] = st.text_area("Fecha de vencimiento CAAF:", "", key="caaf_VENCIMIENTO")
    st.session_state['OBSERVACIONES_EXTRA_CAAF'] = st.text_area("Observaciones extra CAAF:", "", key="caaf_OBSERVACIONES")
    st.header("🔄 Renovación CAA")
    st.session_state['RENOVACION_CAA_STATUS'] = st.selectbox("Estado Renovación CAA:", options=["No aplica", "En curso", "Finalizada", "No iniciada"], key="renovacion_caa_STATUS")

    ren_disabled = st.session_state['RENOVACION_CAA_STATUS'] in ["No aplica", "No iniciada"]
    st.session_state['EXPEDIENTE_RENOVACION_CAA'] = st.text_input("Expediente Renovación CAA:", "N/A", disabled=ren_disabled, key="renovacion_caa_EXPEDIENTE")
    st.session_state['ESTADO_PORTAL_RENOVACION_CAA'] = st.text_input("Estado portal renovación:", "N/A", disabled=ren_disabled, key="renovacion_caa_ESTADO_PORTAL")
    st.session_state['DISPO_RENOVACION_CAA'] = st.text_input("Disposición Renovación:", "N/A", disabled=ren_disabled, key="renovacion_caa_DISPO")
    st.session_state['FECHA_RENOVACION_CAA'] = st.text_input("Fecha de obtencion renovación:", "N/A", disabled=ren_disabled, key="renovacion_caa_FECHA")
    st.session_state['OBSERVACIONES_RENOVACION_CAA'] = st.text_area("Observaciones de renovacion del CAA:", "", key="caa_renov_OBSERVACIONES")
    # --- BLOQUE ÚNICO: HISTORIAL ÚLTIMO CAA ---
    st.header("📜 Historial: Último CAA")
    
    st.selectbox("Estado Último CAA:", options=["Tiene ultimo caa", "No tiene ultimo caa"], key="widget_estado_historial_caa")

    # Esta variable define si se bloquean los campos
    u_caa_bloqueo = st.session_state.get('ULTIMO_CAA_STATUS', '') == "no tiene ultimo caa"

    st.session_state['FECHA_OBTENCION_ULTIMO_CAA'] = st.text_input( "Fecha de obtención último CAA:", value="dd/mm/aaaa", disabled=u_caa_bloqueo, key="input_hist_fecha_caa_u")
    st.session_state['EXPEDIENTE_ULTIMO_CAA'] = st.text_input("Expediente último CAA:",  value="N/A", disabled=u_caa_bloqueo, key="input_hist_exp_caa_u" )
    st.session_state['DISPO_ULTIMO_CAA'] = st.text_input("Disposición último CAA:", value="N/A", disabled=u_caa_bloqueo, key="input_hist_dispo_caa_u")
    st.session_state['OBSERVACIONES_ULTIMO_CAA'] = st.text_area( "Observaciones último CAA:",  value="",  key="input_hist_obs_caa_u")

# --- 3. LEGA y Monitoreos Ambientales ---
with tabs[2]:
    st.header("📜 LEGA")
    st.selectbox("Estado LEGA:", options=["Seleccione...", "vigente", "en curso", "vencida"], key="lega_STATUS")

    lega_disabled = st.session_state['lega_STATUS'] == 'Seleccione...'
    st.session_state['FECHA_OBTENCION_LEGA'] = st.text_input("Fecha de obtención de la LEGA:", value="N/A", disabled=lega_disabled, key="lega_FECHA")
    st.session_state['EXPEDIENTE_LEGA'] = st.text_input("Expediente de la LEGA:", value="N/A", disabled=lega_disabled, key="lega_EXPEDIENTE")
    st.session_state['DISPO_LEGA_VIGENTE'] = st.text_input("Disposicion de la LEGA:", value="N/A", disabled=lega_disabled, key="lega_DISPOSICION")
    st.session_state['ESTADO_LEGA'] = st.text_input("Estado de la LEGA en el portal:", value="N/A", disabled=lega_disabled, key="lega_ESTADO_PORTAL")
    st.session_state['FECHA_VENCIMIENTO_LEGA_VIGENTE'] = st.text_input("Vencimiento de la LEGA:", value="N/A", disabled=lega_disabled, key="lega_VTO")
    st.session_state['OBSERVACIONES_LEGA'] = st.text_area("Observaciones LEGA:", value="N/A", disabled=lega_disabled, key="lega_OBSERVACIONES")
    
# --- SECCIÓN: HISTORIAL ÚLTIMA LEGA ---
    st.header("💨 Historial: Última LEGA")
    
    st.session_state['ULTIMA_LEGA_STATUS'] = st.selectbox("Estado Última LEGA:", options=["Tiene ultima lega", "No tiene ultima LEGA/PDGE", "Tiene permiso PDEG","La ultima lega es la vigente"], key="widget_historial_pdeg_status")

    # Lógica de bloqueo alineada
    u_lega_bloqueo = st.session_state.get('widget_historial_lega_status') == "La ultima lega es la vigente"

    st.session_state['FECHA_ULTIMA_LEGA_PDEG'] = st.text_input( "Fecha obtención última LEGA/PDEG:", value="dd/mm/aaaa", disabled=u_lega_bloqueo, key="key_hist_lega_fecha_u")
    st.session_state['VENCIMIENTO_ULTIMA_LEGA_PDEG'] = st.text_input( "Vencimiento última LEGA/PDEG:", value="dd/mm/aaaa", disabled=u_lega_bloqueo, key="key_hist_lega_vto_u")
    st.session_state['EXPEDIENTE_ULTIMA_LEGA_PDEG'] = st.text_input("Expediente última LEGA/PDEG:", value="N/A",  disabled=u_lega_bloqueo, key="key_hist_lega_exp_u")
    st.session_state['DISPO_ULTIMA_LEGA_PDEG'] = st.text_input( "Disposición última LEGA/PDEG:", value="N/A", disabled=u_lega_bloqueo, key="key_hist_lega_dispo_u")
    st.divider()

  
    st.header("📊 Plan de Monitoreo (NO COMPLETAR ESTE CUADRO)")
    st.write("Recurso | Organismo | Puntos | Parámetros | Frecuencia")

    for i, row_data in enumerate(st.session_state['muestreo_filas_datos']):
        cols = st.columns(5)
        with cols[0]:
            st.session_state['muestreo_filas_datos'][i]['recurso'] = st.text_input("Recurso", value=row_data.get('recurso', ''), key=f"muestreo_recurso_{i}", label_visibility="collapsed")
        with cols[1]:
            st.session_state['muestreo_filas_datos'][i]['organismo'] = st.text_input("Organismo", value=row_data['organismo'], key=f"muestreo_organismo_{i}", label_visibility="collapsed")
        with cols[2]:
            st.session_state['muestreo_filas_datos'][i]['puntos'] = st.text_input("Puntos", value=row_data['puntos'], key=f"muestreo_puntos_{i}", label_visibility="collapsed")
        with cols[3]:
            st.session_state['muestreo_filas_datos'][i]['parametros'] = st.text_input("Parámetros", value=row_data['parametros'], key=f"muestreo_parametros_{i}", label_visibility="collapsed")
        with cols[4]:
            st.session_state['muestreo_filas_datos'][i]['frecuencia'] = st.selectbox("Frecuencia", options=FRECUENCIA_MUESTREO_OPTIONS, index=FRECUENCIA_MUESTREO_OPTIONS.index(row_data['frecuencia']), key=f"muestreo_frecuencia_{i}", label_visibility="collapsed")

    if st.button("Añadir Fila de Monitoreo"):
        st.session_state['muestreo_filas_datos'].append({
            'recurso': '', 'organismo': '', 'puntos': '', 'parametros': '', 'frecuencia': FRECUENCIA_MUESTREO_OPTIONS[0]
        })
        st.rerun()
    st.header("📊 Redactar sobre los monitoreos")

    # Usamos text_area para que el usuario pueda escribir varios renglones
    st.session_state['MONITOREOS_COMENTARIOS'] = st.text_area("Información sobre últimos monitoreos realizados y frecuencia establecida:", value=st.session_state.get('MONITOREOS_COMENTARIOS', "N/A"),height=150,help="Indique aquí fechas de últimos protocolos y periodicidad de los mismos."
)
# --- 4. Autoridad del Agua ---
with tabs[3]:
    st.header("PREFACTIBILIDAD")
    st.session_state['ADA_STATUS'] = st.selectbox("Estado ADA:", options=ADA_DETALLE_ESTADOS, key="ada_STATUS")

    ada_disabled = st.session_state['ADA_STATUS'] == 'Seleccione...' or st.session_state['ADA_STATUS'] == 'No solicitada'
    st.session_state['FECHA_PREFA'] = st.text_input("Fecha de obtención de Prefactibilidad:", "N/A", disabled=ada_disabled, key="ada_FECHA")
    st.session_state['EXPEDIENTE_PREFA'] = st.text_input("Expediente de Prefactibilidad:", "N/A", disabled=ada_disabled, key="ada_EXP")
    st.session_state['VTO_PREFACTIBILIDAD'] = st.text_input("Vencimiento de Prefactibilidad:", "N/A", disabled=ada_disabled, key="ada_VTO")
    st.session_state['NCHI_HIDRAULICA'] = st.text_input("CHi Hidráulica:", "0/1/2/3", disabled=ada_disabled, key="ada_CHI_HID")
    st.session_state['NCHI_EXPLOTACION'] = st.text_input("CHi Explotación:", "0/1/2/3", disabled=ada_disabled, key="ada_CHI_EXP")
    st.session_state['NCHI_VUELCO'] = st.text_input("CHi Vuelco:", "0/1/2/3", disabled=ada_disabled, key="ada_CHI_VUE")
    st.session_state['OBSERVACIONES_PREFA'] = st.text_input("Observaciones Prefactibilidad:", "...", disabled=ada_disabled, key="ada_obs_prefa")
    st.header("HIDRAULICA")
    st.selectbox("Estado Hidraulica:", options=HIDRAULICA_DETALLE_ESTADOS, key="HIDRAULICA_STATUS")
    st.session_state['FECHA_PERMISO_HIDRAULICA_VIGENTE'] = st.text_input("Fecha de obtencion Hidraulica:", key="HIDRAULICA_FECHA")
    st.session_state['RESOL_CONSTANCIA_HIDRAULICA_VIGENTE'] = st.text_input("Resolucion Hidraulica:", key="HIDRAULICA_RESOLUCION")
    st.session_state['EXPEDIENTE_CONSTANCIA_HIDRÁULICA_VIGENTE'] = st.text_input("Expediente Hidraulica:", key="HIDRAULICA_EXPEDIENTE")
    st.session_state['ESTADO_HIDRAULICA'] = st.text_input("Estado de Hidraulica en el portal:", key="HIDRAULICA_ESTADO_PORTAL")
    st.selectbox("tiene ultimo permiso de hidraulica:", options=["tiene ultimo permiso hidraulica", "no tiene ultimo permiso de hidraulica"],  key="ULTIMOHIDRAULICA_STATUS")
    st.session_state['FECHA_CONSTANCIA_HIDRAULICA_OBTENIDA'] = st.text_input("Fecha de obtencion ult. permiso Hidraulica:", "...", key="ada_fecha_ult_hidra")
    st.session_state['RESOL_CONSTANCIA_HIDRAULICA_OBTENIDA'] = st.text_input("Resolucion ult. permiso Hidraulica:", "...",  key="ada_resoc_ult_hidra")
    st.session_state['EXPEDIENTE_CONSTANCIA_HIDRAULICA_OBTENIDA'] = st.text_input("Expediente ult. permiso Hidraulica:", "...", key="ada_exp_ult_hidra")
    st.session_state['OBSERVACIONES_HIDRAULICA'] = st.text_input("Observaciones Hidraulica:", "...", key="ada_obs_hidraulica")
    st.session_state['FECHA_VENCIMIENTO_CONSTANCIA_HIDRAULICA_OBTENIDO'] = st.text_input("Vencimiento ult. permiso Hidraulica:", "...", key="ada_hidra_ult_vto")
    st.header("VUELCO")
    st.selectbox("Estado Vuelco:", options=VUELCO_DETALLE_ESTADOS, key="VUELCO_STATUS")
    st.session_state['FECHA_PERMISO_VUELCO_VIGENTE'] = st.text_input("Fecha de obtencion Vuelco:", key="VUELCO_FECHA")
    st.session_state['RESOL_CONSTANCIA_VUELCO_VIGENTE'] = st.text_input("Resolucion Vuelco:", key="VUELCO_RESOLUCION")
    st.session_state['EXPEDIENTE_CONSTANCIA_VUELCO_VIGENTE'] = st.text_input("Expediente Vuelco:", key="VUELCO_EXPEDIENTE")
    st.session_state['ESTADO_VUELCO'] = st.text_input("Estado de Vuelco en el portal:", key="VUELCO_ESTADO_PORTAL")
    st.selectbox("tiene ultimo permiso de vuelco:", options=["tiene ultimo permiso vuelco", "no tiene ultimo permiso de vuelco"],  key="ULTIMOVUELCO_STATUS")
    st.session_state['FECHA_PERMISO_VUELCO_OBTENIDO'] = st.text_input("Fecha de obtencion ult. permiso Vuelco:", "...", key="ada_fecha_ult_vuelco")
    st.session_state['RESOL_PERMISO_VUELCO_OBTENIDO'] = st.text_input("Resolucion ult. permiso Vuelco:", "...", key="ada_resoc_ult_vuelco")
    st.session_state['EXPEDIENTE_PERMISO_VUELCO_OBTENIDO'] = st.text_input("Expediente ult. permiso Vuelco:", "...", key="ada_exp_ult_vuelco")
    st.session_state['FECHA_VENCIMIENTO_PERMISO_VUELCO_OBTENIDO'] = st.text_input("Vencimiento ult. permiso Vuelco:", "...", key="ada_vuelco_ult_vto")
    st.session_state['OBSERVACIONES_VUELCO'] = st.text_input("Observaciones Vuelco","...", key="ada_obs_vuelco")
    st.header("EXPLOTACION")
    st.selectbox("Estado Explotacion:", options=EXPLOTACION_DETALLE_ESTADOS, key="EXPLOTACION_STATUS")
    st.session_state['FECHA_PERMISO_EXPLOTACION_VIGENTE'] = st.text_input("Fecha de obtencion Explotacion:", key="EXPLOTACION_FECHA")
    st.session_state['RESOL_CONSTANCIA_EXPLOTACION_VIGENTE'] = st.text_input("Resolucion Explotacion:", key="EXPLOTACION_RESOLUCION")
    st.session_state['EXPEDIENTE_CONSTANCIA_EXPLOTACION_VIGENTE'] = st.text_input("Expediente Explotacion:", key="EXPLOTACION_EXPEDIENTE")
    st.session_state['ESTADO_EXPLOTACION'] = st.text_input("Estado de Explotacion en el portal:", key="EXPLOTACION_ESTADO_PORTAL")
    st.selectbox("tiene ultimo permiso de hidráulica:", options=["tiene ultimo permiso explotacion", "no tiene ultimo permiso de explotacion"],  key="ULTIMOEXPLOTACION_STATUS")
    st.session_state['FECHA_PERMISO_EXPLOTACION_OBTENIDO'] = st.text_input("Fecha de obtencion ult. permiso Explotacion:", "...", key="ada_fecha_ult_exp")
    st.session_state['RESOL_PERMISO_EXPLOTACION_OBTENIDO'] = st.text_input("Resolucion ult. permiso Explotacion:", "...", key="ada_resoc_ult_exp")
    st.session_state['EXPEDIENTE_PERMISO_EXPLOTACION_OBTENIDO'] = st.text_input("Vencimiento ult. permiso Explotacion:", "...", key="ada_exp_ult_vto")
    st.session_state['FECHA_VENCIMIENTO_PERMISO_EXPLOTACION_OBTENIDO'] = st.text_input("Expediente ult. permiso Explotacion:", "...", key="ada_exp_ult_exp")
    st.session_state['OBSERVACIONES_EXPLOTACION'] = st.text_input("Observaciones Explotacion:", "...", key="ada_obs_explotacion")
    st.header("RED DE MONITOREO")
    st.text_input("Redactar sobre Red de monitoreo:", key="RED_MONITOREOS")
    st.header("TASAS Y CANON")
    st.session_state['COMPLETAR_TASA_CANON_ADA'] = st.text_input("Redactar sobre el estado de Tasa y Canon:", key="TASA_CANON")
# --- 5. Otros requisitos de Agua' ---
with tabs[4]:  
    st.header("🗺️ ACUMAR (Autoridad de Cuenca Matanza Riachuelo)")
    st.session_state['ACUMAR_STATUS'] = st.selectbox("Estado ACUMAR:", options=ACUMAR_DETALLE_ESTADOS, key="acumar_STATUS")

    acumar_disabled = st.session_state['ACUMAR_STATUS'] == 'Seleccione...' or st.session_state['ACUMAR_STATUS'] == 'No aplica' or st.session_state['ACUMAR_STATUS'] == 'no aplica'
    st.session_state['ACUMAR_NIA'] = st.text_input("NIA ACUMAR:", "N/A", disabled=acumar_disabled, key="acumar_NIA")
    st.session_state['ACUMAR_AÑO_DDJJ'] = st.text_input("AÑO DDJJ ACUMAR:", "N/A", disabled=acumar_disabled, key="acumar_AÑO")
    st.session_state['ACUMAR_OBSERVACIONES'] = st.text_area("Observaciones ACUMAR:", "", disabled=acumar_disabled, key="acumar_OBS")
    st.selectbox("Estado ACUMAR:", options=["Plan de adecuacion en curso","No aplica Plan de adecuacion"], key="PLANDEADECUACIONACUMAR_STATUS")
    st.header("🗺️ AySA y DPyRA")
    st.session_state['AYSA_STATUS'] = st.selectbox("Estado AySA:", options=AYSA_DETALLE_ESTADOS, key="aysa_STATUS")
    st.session_state['REDACCION_AYSA'] = st.text_input("Redactar sobre la situacion de la empresa respecto a Aysa:", "N/A", key="se_AYSA_AP")
# --- 5. Otros Organismos ---
with tabs[5]:

    st.header("⚡ Secretaría de Energía")
   
    st.subheader("Inscripción 1102/04")
    st.session_state['INSCRIPCION_1102'] = st.selectbox("Inscripción 1102/04:", options=["Inscripto", "No inscripto", "No aplica"], key="se_1102_STATUS")
    st.session_state['NUMERO_SE'] = st.text_input("Número de operador (1102/04):", "N/A", key="se_NUM_SE")
    st.session_state['SE_EXPEDIENTE'] = st.text_input("Expediente Sec. Energía:", "N/A", key="se_EXP")
    st.session_state['OBSERVACIONES_1102'] = st.text_area("Observaciones Res. 1102:", "", key="se_OBS_1102")
    st.subheader("Auditoría de seguridad 404/94")
    st.session_state['AUDITORIA_404'] = st.selectbox("Auditoría de seguridad 404/94:", options=["Realizo", "No realizo", "No aplica", "No inscripto, no realiza"], key="se_404_STATUS")
    auditoria_404_disabled = st.session_state['AUDITORIA_404'] not in ["Realizo", "No realizo"]
    st.session_state['CERTIFICADO_TANQUES_AEREOS_VTO'] = st.text_input("Vencimiento de tanques aereos:", "N/A", disabled=auditoria_404_disabled, key="se_VTO_AEREOS")
    st.session_state['CANTIDAD_DE_TANQUES_AEREOS'] = st.text_input("N.º tanques aereos:", "N/A", disabled=auditoria_404_disabled, key="se_N_AEREOS")
    st.session_state['CERTIFICADO_TANQUES_SUB_VTO'] = st.text_input("Vencimiento de tanques subterraneos:", "N/A", disabled=auditoria_404_disabled, key="se_VTO_SUBTERRANEOS")
    st.session_state['CANTIDAD_DE_TANQUES_SUB'] = st.text_input("N.º tanques subterraneos:", "N/A", disabled=auditoria_404_disabled, key="se_N_SUBTERRANEOS")
    st.session_state['SE_CERTIFICADO_SUB_HERMETICIDAD'] = st.text_input("Vencimiento de hermeticidad subterraneos:", "N/A", disabled=auditoria_404_disabled, key="se_VTO_SUBTERRANEOS_HERMETICIDAD")
    st.session_state['SE_CERTIFICADO_TAMBORES_VENCIMIENTO'] = st.text_input("Vencimiento de tambores de envase:", "N/A", disabled=auditoria_404_disabled, key="se_VTO_TAMBORES")
    st.session_state['OBSERVACIONES_404'] = st.text_area("Observaciones Res. 404/94:", "", key="se_OBS_404")
    st.subheader("Inscripción 277/25")
    st.session_state['INSCRIPCION_277'] = st.selectbox("Inscripción 277/25:", options=["Aplica", "No aplica"], key="se_277_STATUS")
    st.session_state['AUDITORIA277_VENCIMIENTO'] = st.text_area("Vencimiento 277/25:", "", key="se_277_VTO")
    st.session_state['OBSERVACIONES_277'] = st.text_area("Observaciones 277/25:", "", key="se_OBS_277")

   


    st.header("🚧 RENPRE y Seguro Ambiental")
    st.session_state['RENPRE_STATUS'] = st.selectbox("Estado RENPRE:", options=RENPRE_DETALLE_ESTADOS, key="renpre_STATUS")
    renpre_disabled = st.session_state['RENPRE_STATUS'] == 'Seleccione...' or st.session_state['RENPRE_STATUS'] == 'No aplica'
    st.session_state['NUMERO_RENPRE'] = st.text_input("Número de operador RENPRE:", "N/A", disabled=renpre_disabled, key="renpre_NUM")
    st.session_state['VENCIMIENTO_RENPRE'] = st.text_input("Fecha de venciemiento RENPRE:", "N/A", disabled=renpre_disabled, key="renpre_VTO")

    st.markdown("--- --- --- --- --- --- --- --- --- --- --- --- --- --- --- --- --- --- --- ---")
    st.session_state['SEGURO_STATUS'] = st.selectbox("Seguro ambiental:", options=["Vigente", "Vencida", "Nunca Tuvo"], key="seguro_STATUS")
    seguro_disabled = st.session_state['SEGURO_STATUS'] == 'Nunca Tuvo'
    st.session_state['NUMERO_POLIZA'] = st.text_input("Nº Póliza:", "N/A", disabled=seguro_disabled, key="seguro_POLIZA_NUM")
    st.session_state['VTO_POLIZA'] = st.text_input("Vencimiento de póliza:", "N/A", disabled=seguro_disabled, key="seguro_POLIZA_VTO")

# --- 6. Residuos y ASP ---
with tabs[6]:
    st.header("1. Residuos Especiales (Inscripción)")
    st.session_state['RESIDUOS_ESPECIALES_STATUS'] = st.selectbox("Estado Generador (General):", options=RESIDUOS_ESPECIALES_STATUS_GENERAL, key="rree_STATUS")

    st.header("2. Certificado de Habilitación Especial (CHE)")
    st.session_state['CHE_STATUS'] = st.selectbox("Estado CHE:", options=CHE_DETALLE_ESTADOS, key="rree_CHE_STATUS")

    che_disabled = st.session_state['CHE_STATUS'] == 'Seleccione...'
    st.session_state['ANIO_CHE'] = st.text_input("Año de obtención del CHE:", "2025", disabled=che_disabled, key="rree_ANIO_CHE")
    st.session_state['OBSERVACIONES_TICKETS_CONSULTA_CHE'] = st.text_area("Observaciones CHE:", "", disabled=che_disabled, key="rree_OBS_CHE")

    st.header("♻️ 3. Gestión Operativa de Residuos")
    st.text_input("Situacion de gestion de residuos(redactar):", "", key="REDACCION_RESIDUOS")
    st.session_state['GIRSU_STATUS'] = st.selectbox("GIRSU:", options=["Aplica", "No aplica","Aplica y realizo la presentacion", "Aplica y no realizo la presentacion"], key="rree_GIRSU_STATUS")
    st.session_state['FECHA_PRESENTACION_GIRSU'] = st.text_area("Fecha de presentacion GIRSU:", "", key="rree_GIRSU_FECHA")
    st.session_state['PATOGENICOS_STATUS'] = st.selectbox("Estado Residuos patogénicos:", options=["Inscripto", "No inscripto", "No aplica"], key="rree_PATOGENICOS_STATUS")
    st.session_state['PATOGENICOS_EXPEDIENTE'] = st.text_area("Expediente Patogenicos:", "",  key="PATOGENICOS_EXP")
    st.session_state['PATOGENICOS_FECHA'] = st.text_area("Fecha de inscripcion Patogenicos:", "",  key="PATO_FECHA")
    st.header("🌡️ ASP (Aparatos a Presión)")
    st.session_state['ASP_STATUS'] = st.selectbox("Estado ASP:", options=["Seleccione...", "Finalizada", "Caratulada", "No Presentado"], key="asp_STATUS")

    asp_disabled = st.session_state['ASP_STATUS'] == 'Seleccione...' or st.session_state['ASP_STATUS'] == 'No Presentado'
    st.session_state['VENCIMIENTO_ASP'] = st.text_input("Vencimiento de presentación de ASP:", "N/A", disabled=asp_disabled, key="asp_VTO")
    st.session_state['EXPEDIENTE_ASP'] = st.text_input("Expediente del ASP:", "N/A", disabled=asp_disabled, key="asp_EXPEDIENTE")
    st.session_state['OBSERVACIONES_EXTRA_ASP'] = st.text_area("Observaciones ASP:", "", disabled=asp_disabled, key="asp_OBS")

    st.subheader("Calibración de Válvulas de Seguridad")
    st.session_state['VALVULAS_STATUS'] = st.selectbox("Calibración de válvulas de seguridad:", options=["Cumple", "No Cumple"], key="asp_VALVULAS_STATUS")
    valvulas_disabled = st.session_state['VALVULAS_STATUS'] == 'No Cumple'
    st.session_state['VENCIMIENTO_CALIBRACION_ASP'] = st.text_input("Vencimiento de válvulas:", "N/A", disabled=valvulas_disabled, key="asp_VTO_VALVULAS")


# --- 7. Hallazgos y Recomendaciones ---
with tabs[7]:
    st.header("🔍 Hallazgos de Campo")

    # 1. Función de autocompletado corregida y blindada
    def aplicar_preset_hallazgo(idx):
        # Obtenemos la versión actual para construir la llave correcta
        ver_actual = st.session_state.get(f"version_{idx}", 0)
        clave_preset = f"hallazgo_preset_{idx}_{ver_actual}"
        seleccion = st.session_state.get(clave_preset)
        
        if seleccion and seleccion != "Autocompletar (Seleccione)...":
            for cat, items in HALLAZGOS_PREDEFINIDOS.items():
                for item in items:
                    # Coincidencia exacta con el formato del selector
                    texto_opcion = f"[{cat}] {item.get('situacion', '')[:80]}"
                    
                    if texto_opcion == seleccion:
                        # Actualizamos los datos en el diccionario de la lista
                        st.session_state['hallazgos_widgets_list'][idx].update({
                            'situacion': item.get('situacion', ''),
                            'autoridad': item.get('autoridad', ''),
                            'riesgo': item.get('riesgo', ''),
                            'recomendacion': item.get('recomendacion', '')
                        })
                        
                        # Incrementamos la versión para "forzar" a Streamlit a refrescar los campos
                        st.session_state[f"version_{idx}"] = ver_actual + 1
                        return

    # 2. Bucle de visualización de hallazgos
    # Usamos list(range(...)) para evitar errores de mutación durante el renderizado
    hallazgos_actuales = st.session_state.get('hallazgos_widgets_list', [])
    
    for i in range(len(hallazgos_actuales)):
        finding = hallazgos_actuales[i]
        ver = st.session_state.get(f"version_{i}", 0)
        
        # Expander con ID único para evitar colisiones
        with st.expander(f"📌 Hallazgo # {i+1}", expanded=True):
            
            # Construcción de opciones de autocompletado
            opciones_preset = ["Autocompletar (Seleccione)..."]
            for cat, lista in HALLAZGOS_PREDEFINIDOS.items():
                for h in lista:
                    opciones_preset.append(f"[{cat}] {h.get('situacion', '')[:80]}")

            # Selector de Preset con Key dinámica por versión
            st.selectbox(
                "Seleccionar un modelo predefinido:",
                options=opciones_preset,
                key=f"hallazgo_preset_{i}_{ver}", 
                on_change=aplicar_preset_hallazgo,
                args=(i,)
            )

            st.divider()
            
            # Formulario de edición con keys únicas basadas en índice y versión
            st.session_state['hallazgos_widgets_list'][i]['situacion'] = st.text_area(
                "Situación:", value=finding['situacion'], key=f"h_sit_{i}_{ver}", height=150)
            
            st.session_state['hallazgos_widgets_list'][i]['autoridad'] = st.text_input(
                "Autoridad:", value=finding['autoridad'], key=f"h_aut_{i}_{ver}")
            
            st.session_state['hallazgos_widgets_list'][i]['riesgo'] = st.text_area(
                "Riesgo:", value=finding['riesgo'], key=f"h_rie_{i}_{ver}", height=100)
            
            st.session_state['hallazgos_widgets_list'][i]['recomendacion'] = st.text_area(
                "Recomendación:", value=finding['recomendacion'], key=f"h_rec_{i}_{ver}", height=100)

            # Botón para eliminar con key única
            if st.button(f"🗑️ Quitar Hallazgo {i+1}", key=f"del_btn_{i}_{ver}"):
                st.session_state['hallazgos_widgets_list'].pop(i)
                # Al eliminar, reseteamos las versiones para evitar desajustes de índice
                for j in range(len(st.session_state['hallazgos_widgets_list']) + 1):
                    if f"version_{j}" in st.session_state:
                        st.session_state[f"version_{j}"] = 0
                st.rerun()

    # Botón para añadir nuevo hallazgo
    if st.button("➕ Añadir Hallazgo Nuevo", key="add_finding_main"):
        st.session_state['hallazgos_widgets_list'].append({
            'situacion': '', 'autoridad': '', 'riesgo': '', 'recomendacion': ''
        })
        st.rerun()
# --- GENERATE REPORT BUTTON ---
st.markdown("--- --- --- --- --- --- --- --- --- --- --- --- --- --- --- --- --- --- --- ---")
if st.session_state['uploaded_file'] is None:
    st.error("Por favor, sube una plantilla DOCX en la barra lateral para generar el informe.")
else:
    if st.button("GENERAR INFORME", use_container_width=True, type="primary"):
        try:
            # Load the document from the uploaded file in memory
            doc_stream = io.BytesIO(st.session_state['uploaded_file'].getvalue())
            doc = Document(doc_stream)

            # 1. Collect all data from st.session_state (Variables simples)
            user_data = {
                'RAZON_SOCIAL': st.session_state.get('main_NOMBRE_EMPRESA', ''),
                'NOMBRE_EMPRESA': st.session_state.get('main_NOMBRE_EMPRESA', ''),
                'NOMBRE_PLANTA': st.session_state.get('main_NOMBRE_PLANTA', ''),
                'MES_AUDITORIA': st.session_state.get('main_MES_AUDITORIA', ''),
                'DIRECCION_EMPRESA': st.session_state.get('main_DIRECCION', ''),
                'MUNICIPIO_EMPRESA': st.session_state.get('main_MUNICIPIO', ''),
                'RUBRO_EMPRESA': st.session_state.get('main_RUBRO', ''),

                'FECHA_HABILITACION': st.session_state.get('hab_FECHA', ''),
                'EXPEDIENTE_HABILITACION': st.session_state.get('hab_EXPEDIENTE', ''),
                'OBSERVACION_HAB_MUNICIPAL': st.session_state.get('hab_OBSERVACION', ''),
                
                'FECHA_CNCA': st.session_state.get('cnca_FECHA', ''),
                'EXPEDIENTE_CNCA': st.session_state.get('cnca_EXPEDIENTE', ''),
                'CATEGORIA_CNCA': st.session_state.get('cnca_CATEGORIA', ''),
                'PUNTOS_CNCA': st.session_state.get('cnca_PUNTOS', ''),
                'VENCIMIENTO_CNCA': st.session_state.get('cnca_VENCIMIENTO', ''),
                'DISPO_CNCA': st.session_state.get('cnca_DISPO', ''),
                'OBSERVACIONES_CNCA': st.session_state.get('cnca_OBSERVACIONES', ''),

      
                'FECHA_OBTENCION_CAAP': st.session_state.get('caap_FECHA', ''),
                'EXPEDIENTE_CAAP': st.session_state.get('caap_EXP', ''),
                'FECHA_CAAP': st.session_state.get('caap_FECHA',''),
                'EXP_CAAP': st.session_state.get('caap_EXP',''),
                'VTO_CAAP': st.session_state.get('caap_VTO',''),
                'DISPO_CAAP': st.session_state.get('caap_DISPO', ''),
                'VIGENCIA_CAAP': st.session_state.get('caap_VIGENCIA', ''),
                'VENCIMIENTO_CAAP': st.session_state.get('caap_VTO', ''),
                'ESTADO_PORTAL_CAAP': st.session_state.get('caap_ESTADO_PORTAL', ''),
                'OBSERVACIONES_CAAP': st.session_state.get('caap_OBSERVACIONES', ''),

                'FECHA_CAAF': st.session_state.get('caaf_FECHA', ''),
                'EXPEDIENTE_CAAF': st.session_state.get('caaf_EXPEDIENTE', ''),
                'DISPO_CAAF': st.session_state.get('caaf_DISPO', ''),
                'ESTADO_CAAF': st.session_state.get('caaf_ESTADO_PORTAL', ''),
                'VENCIMIENTO_CAAF': st.session_state.get('caaf_VENCIMIENTO', ''),
                'OBSERVACIONES_CAAF': st.session_state.get('caaf_OBSERVACIONES', ''),


                'ESTADO_RENOVACION_CAA': st.session_state.get('renovacion_caa_STATUS', ''),
                'EXPEDIENTE_RENOVACION_CAA': st.session_state.get('renovacion_caa_EXPEDIENTE', ''),
                'ESTADO_PORTAL_RENOVACION_CAA': st.session_state.get('renovacion_caa_ESTADO_PORTAL', ''),
                'DISPO_RENOVACION_CAA': st.session_state.get('renovacion_caa_DISPO', ''),
            
                # --- HISTORIAL ÚLTIMO CAA ---
                'FECHA_ULTIMO_CAA': st.session_state.get('input_hist_fecha_caa_u', ''),
                'EXPEDIENTE_ULTIMO_CAA': st.session_state.get('input_hist_exp_caa_u', ''),
                'DISPO_ULTIMO_CAA': st.session_state.get('input_hist_dispo_caa_u', ''),
                'OBSERVACIONES_ULTIMO_CAA': st.session_state.get('input_hist_obs_caa_u', ''), 


                # --- HISTORIAL ÚLTIMO CAA ---
                'FECHA_OBTENCION_ULTIMO_CAA': st.session_state.get('input_hist_fecha_caa_u', ''),
                'EXPEDIENTE_ULTIMO_CAA': st.session_state.get('input_hist_exp_caa_u', ''),
                'DISPO_ULTIMO_CAA': st.session_state.get('input_hist_dispo_caa_u', ''),
                'OBSERVACIONES_ULTIMO_CAA': st.session_state.get('ultimo_caa_OBSERVACIONES', ''),
                

                'ESTADO_RENOVACION_CAA': st.session_state.get('renovacion_caa_STATUS', ''),
                'EXPEDIENTE_RENOVACION_CAA': st.session_state.get('renovacion_caa_EXPEDIENTE', ''),
                'ESTADO_PORTAL_RENOVACION_CAA': st.session_state.get('renovacion_caa_ESTADO_PORTAL', ''),
                'DISPO_RENOVACION_CAA': st.session_state.get('renovacion_caa_DISPO', ''),
                'FECHA_RENOVACION_CAA': st.session_state.get ('renovacion_caa_FECHA',''),
                'OBSERVACIONES_RENOVACION_CAA': st.session_state.get('caa_renov_OBSERVACIONES',''),
                'FECHA_OBTENCION_LEGA': st.session_state.get('lega_FECHA', ''),
                'EXPEDIENTE_LEGA': st.session_state.get('lega_EXPEDIENTE', ''),
                'ESTADO_LEGA': st.session_state.get('lega_ESTADO_PORTAL', ''),
                'DISPO_LEGA_VIGENTE': st.session_state.get('lega_DISPOSICION',''),
                'FECHA_VENCIMIENTO_LEGA_VIGENTE': st.session_state.get('lega_VTO', ''),
                'OBSERVACIONES_LEGA': st.session_state.get('lega_OBSERVACIONES', ''),
                'MONITOREOS_COMENTARIOS': st.session_state.get('MONITOREOS_COMENTARIOS', ''),
                # --- HISTORIAL ÚLTIMA LEGA (GASEOSOS) ---
                'DISPO_ULTIMA_LEGA': st.session_state.get('ultima_lega_DISPO', ''),
                'OBSERVACIONES_ULTIMA_LEGA': st.session_state.get('ultima_lega_OBSERVACIONES', ''),
                'FECHA_OBTENCION_LEGA': st.session_state.get('lega_FECHA', ''),
                'EXPEDIENTE_LEGA': st.session_state.get('lega_EXPEDIENTE', ''),
                'VENCIMIENTO_LEGA': st.session_state.get('lega_VTO', ''),
                'FECHA_ULTIMA_LEGA': st.session_state.get('key_hist_lega_fecha_u', ''),
                'VENCIMIENTO_ULTIMA_LEGA': st.session_state.get('key_hist_lega_vto_u', ''),
                'EXPEDIENTE_ULTIMA_LEGA': st.session_state.get('key_hist_lega_exp_u', ''),
                
                'RESIDUOS_ESPECIALES_STATUS': st.session_state.get('RESIDUOS_ESPECIALES_STATUS', ''),
                'CHE_STATUS': st.session_state.get('CHE_STATUS', ''),
                'ANIO_CHE': st.session_state.get('ANIO_CHE', ''),
                'OBSERVACIONES_TICKETS_CONSULTA_CHE': st.session_state.get('OBSERVACIONES_TICKETS_CONSULTA_CHE', ''),
                'GESTION_RESIDUOS_STATUS': st.session_state.get('GESTION_RESIDUOS_STATUS', ''),
                'TIPO_RESIDUO': st.session_state.get('TIPO_RESIDUO', ''),
                'OBSERVACION_EXTRA_RESIDUOS': st.session_state.get('OBSERVACION_EXTRA_RESIDUOS', ''),
                'GIRSU_STATUS': st.session_state.get('GIRSU_STATUS', ''),
                'FECHA_PRESENTACION_GIRSU':st.session_state.get('FECHA_PRESENTACION_GIRSU',''),
                'PATOGENICOS_STATUS': st.session_state.get('PATOGENICOS_STATUS', ''),
                'PATOGENICOS_EXPEDIENTE' : st.session_state.get('PATOGENICOS_EXP', ''),
                'PATOGENICOS_FECHA' : st.session_state.get('PATO_FECHA',''),
                'REDACCION_RESIDUOS': st.session_state.get('REDACCION_RESIDUOS', ''),

                'ASP_STATUS': st.session_state.get('ASP_STATUS', ''),
                'VENCIMIENTO_ASP': st.session_state.get('VENCIMIENTO_ASP', ''),
                'EXPEDIENTE_ASP': st.session_state.get('EXPEDIENTE_ASP', ''),
                'OBSERVACIONES_EXTRA_ASP': st.session_state.get('OBSERVACIONES_EXTRA_ASP', ''),
                'VALVULAS_STATUS': st.session_state.get('VALVULAS_STATUS', ''),
                'VENCIMIENTO_CALIBRACION_ASP': st.session_state.get('VENCIMIENTO_CALIBRACION_ASP', ''),

                'FECHA_PREFA': st.session_state.get('FECHA_PREFA', ''),
                'EXPEDIENTE_PREFA': st.session_state.get('EXPEDIENTE_PREFA', ''),
                'NCHI_HIDRAULICA': st.session_state.get('ada_CHI_HID', ''),
                'NCHI_EXPLOTACION': st.session_state.get('ada_CHI_EXP', ''),
                'NCHI_VUELCO': st.session_state.get('ada_CHI_VUE', ''),
                'VTO_PREFACTIBILIDAD': st.session_state.get('ada_VTO',''),
                'ESTADO_PERMISO_HIDRAULICA': st.session_state.get('ESTADO_PERMISO_HIDRAULICA', ''),
                'ESTADO_PERMISO_EXPLOTACION': st.session_state.get('ESTADO_PERMISO_EXPLOTACION', ''),
                'ESTADO_PERMISO_VUELCO': st.session_state.get('ESTADO_PERMISO_VUELCO', ''),
                'RED_MONITOREOS': st.session_state.get('RED_MONITOREOS', ''),
                'FECHA_VENCIMIENTO_CONSTANCIA_HIDRAULICA_OBTENIDO': st.session_state.get('FECHA_VENCIMIENTO_CONSTANCIA_HIDRAULICA_OBTENIDO',''),
                'FECHA_VENCIMIENTO_PERMISO_VUELCO_OBTENIDO': st.session_state.get('FECHA_VENCIMIENTO_PERMISO_VUELCO_OBTENIDO',''),
                'FECHA_VENCIMIENTO_PERMISO_EXPLOTACION_OBTENIDO': st.session_state.get('FECHA_VENCIMIENTO_PERMISO_EXPLOTACION_OBTENIDO',''),
                'RENPRE_STATUS': st.session_state.get('RENPRE_STATUS', ''),
                'NUMERO_RENPRE': st.session_state.get('NUMERO_RENPRE', ''),
                'VENCIMIENTO_RENPRE': st.session_state.get('VENCIMIENTO_RENPRE', ''),
                # EXPLOTACIÓN (Sincronización de nombres)
                'FECHA_PERMISO_EXPLOTACION_OBTENIDO': st.session_state.get('FECHA_PERMISO_EXPLOTACION_OBTENIDO', ''),
                'RESOL_PERMISO_EXPLOTACION_OBTENIDO': st.session_state.get('RESOL_PERMISO_EXPLOTACION_OBTENIDO', ''),
                'EXPEDIENTE_PERMISO_EXPLOTACION_OBTENIDO': st.session_state.get('ada_exp_ult_hidra', ''),
          
                'COMPLETAR_TASA_CANON_ADA':st.session_state.get('COMPLETAR_TASA_CANON_ADA',''),
                # HIDRÁULICA (Sincronización de nombres)
                'FECHA_CONSTANCIA_HIDRAULICA_OBTENIDA': st.session_state.get('FECHA_CONSTANCIA_HIDRAULICA_OBTENIDA', ''),
                'RESOL_CONSTANCIA_HIDRAULICA_OBTENIDA': st.session_state.get('RESOL_CONSTANCIA_HIDRAULICA_OBTENIDA', ''),
                'EXPEDIENTE_CONSTANCIA_HIDRAULICA_OBTENIDA': st.session_state.get('EXPEDIENTE_CONSTANCIA_HIDRAULICA_VIGENTE', ''),

                # VUELCO (Sincronización de nombres)
                'FECHA_PERMISO_VUELCO_OBTENIDO': st.session_state.get('FECHA_PERMISO_VUELCO_OBTENIDO', ''),
                'RESOL_PERMISO_VUELCO_OBTENIDO': st.session_state.get('RESOL_PERMISO_VUELCO_OBTENIDO', ''),
                'EXPEDIENTE_PERMISO_VUELCO_OBTENIDO': st.session_state.get('EXPEDIENTE_PERMISO_VUELCO_OBTENIDO', ''),
                'OBSERVACIONES_HIDRAULICA':st.session_state.get('OBSERVACIONES_HIDRAULICA',''),
                'OBSERVACIONES_VUELCO':st.session_state.get('OBSERVACIONES_VUELCO',''),
                'OBSERVACIONES_EXPLOTACION':st.session_state.get('OBSERVACIONES_EXPLOTACION',''),

                'SEGURO_STATUS': st.session_state.get('SEGURO_STATUS', ''),
                'NUMERO_POLIZA': st.session_state.get('NUMERO_POLIZA', ''),
                'VTO_POLIZA': st.session_state.get('VTO_POLIZA', ''),

                'ACUMAR_STATUS': st.session_state.get('ACUMAR_STATUS', ''),
                'ACUMAR_EXPEDIENTE': st.session_state.get('ACUMAR_EXPEDIENTE', ''),
                'ACUMAR_OBSERVACIONES': st.session_state.get('ACUMAR_OBSERVACIONES', ''),
                'ACUMAR_NIA': st.session_state.get('acumar_NIA',''),
                'ACUMAR_AÑO_DDJJ': st.session_state.get('acumar_AÑO',''),

                'SE_STATUS': st.session_state.get('SE_STATUS', ''),
                'SE_EXPEDIENTE': st.session_state.get('SE_EXPEDIENTE', ''),
                'INSCRIPCION_1102': st.session_state.get('INSCRIPCION_1102', ''),
                'NUMERO_SE': st.session_state.get('NUMERO_SE', ''),
                'AUDITORIA_404': st.session_state.get('AUDITORIA_404', ''),
                'CANTIDAD_DE_TANQUES': st.session_state.get('CANTIDAD_DE_TANQUES', ''),
                'VENCIMIENTO_AUDITORIA404': st.session_state.get('VENCIMIENTO_AUDITORIA404', ''),
                'INSCRIPCION_277': st.session_state.get('INSCRIPCION_277', ''),
                'OBSERVACIONES_277': st.session_state.get('OBSERVACIONES_277', ''),
                'CANTIDAD_DE_TANQUES_AEREOS': st.session_state.get('se_N_AEREOS', 'N/A'),
                'CERTIFICADO_TANQUES_AEREOS_VTO': st.session_state.get('se_VTO_AEREOS', 'N/A'),
                'CANTIDAD_DE_TANQUES_SUB': st.session_state.get('se_N_SUBTERRANEOS', 'N/A'),
                'CERTIFICADO_TANQUES_SUB_VTO': st.session_state.get('se_VTO_SUBTERRANEOS', 'N/A'),
                'SE_CERTIFICADO_SUB_HERMETICIDAD': st.session_state.get('se_VTO_SUBTERRANEOS_HERMETICIDAD', 'N/A'),
                'SE_CERTIFICADO_TAMBORES_VENCIMIENTO': st.session_state.get('se_VTO_TAMBORES', 'N/A'),
                'AUDITORIA277_VENCIMIENTO': st.session_state.get('se_277_VTO', ''),
                'OBSERVACIONES_1102': st.session_state.get('se_OBS_1102',''),
                'OBSERVACIONES_404': st.session_state.get('se_OBS_404',''),
                "CONCLUSIONES_GENERALES": st.session_state.get('CONCLUSIONES_GENERALES', ''),
                # Hidráulica Vigente
                'FECHA_ CONSTANCIA_HIDRAULICA_VIGENTE': st.session_state.get('HIDRAULICA_FECHA', ''),
                'RESOL_CONSTANCIA_HIDRAULICA_VIGENTE': st.session_state.get('HIDRAULICA_RESOLUCION', ''),
                'EXPEDIENTE_CONSTANCIA_HIDRAULICA_VIGENTE': st.session_state.get('HIDRAULICA_EXPEDIENTE', ''),
                'ESTADO_HIDRAULICA': st.session_state.get('HIDRAULICA_ESTADO_PORTAL', ''),

                # Vuelco Vigente
                'FECHA_ PERMISO_VUELCO_VIGENTE': st.session_state.get('VUELCO_FECHA', ''),
                'RESOL_PERMISO_VUELCO_VIGENTE': st.session_state.get('VUELCO_RESOLUCION', ''),
                'EXPEDIENTE_PERMISO_VUELCO_VIGENTE': st.session_state.get('VUELCO_EXPEDIENTE', ''),
                'ESTADO_VUELCO': st.session_state.get('VUELCO_ESTADO_PORTAL', ''),

                # Explotación Vigente
                'FECHA_ PERMISO_EXPLOTACION_VIGENTE': st.session_state.get('EXPLOTACION_FECHA', ''),
                'RESOL_PERMISO_EXPLOTACION_VIGENTE': st.session_state.get('EXPLOTACION_RESOLUCION', ''),
                'EXPEDIENTE_PERMISO_EXPLOTACION_VIGENTE': st.session_state.get('EXPLOTACION_EXPEDIENTE', ''),
                'ESTADO_EXPLOTACION': st.session_state.get('EXPLOTACION_ESTADO_PORTAL', ''),

                # Historial ADA (Campos extras que agregaste)
                'FECHA_ULT_HIDRA': st.session_state.get('ada_fecha_ult_hidra', ''),
                'RESOL_ULT_HIDRA': st.session_state.get('ada_resoc_ult_hidra', ''),
                'FECHA_ULT_VUELCO': st.session_state.get('ada_fecha_ult_vuelco', ''),
                'RESOL_ULT_VUELCO': st.session_state.get('ada_resoc_ult_vuelco', ''),
                'REDACCION_AYSA':st.session_state.get('se_AYSA_AP',''),

            }
            reemplazar_marcadores(doc, user_data)
            # 2. Process Conditional Logic (Remove blocks)
            # Aquí vinculamos el valor del widget con la lógica del diccionario MARCADORES_CONDICIONALES
            
            s_hab = st.session_state.get('hab_STATUS', 'Seleccione...')
            for p in find_paragraphs_to_remove(doc, s_hab, "HABILITACION_MUNICIPAL"): remove_paragraph(p)

            s_cnca = st.session_state.get('cnca_STATUS', 'No aplica')
            for p in find_paragraphs_to_remove(doc, s_cnca, "CNCA_STATUS"): remove_paragraph(p)

            s_caap = st.session_state.get('caap_CAAP_LOGICA_ESTADOS', 'No aplica')
            for p in find_paragraphs_to_remove(doc, s_caap, "CAAP_STATUS"): remove_paragraph(p)

            s_caaf = st.session_state.get('caaf_CAAF_LOGICA', 'No aplica')
            if s_caaf == "No aplica": s_caaf = "Eliminar todo"
            for p in find_paragraphs_to_remove(doc, s_caaf, "CAAF_STATUS"): remove_paragraph(p)

            # --- NUEVOS HISTORIALES ---
            s_u_caa = st.session_state.get('widget_estado_historial_caa', 'No tiene ultimo caa')
            for p in find_paragraphs_to_remove(doc, s_u_caa, "ULTIMO_CAA_STATUS"): remove_paragraph(p)

            s_u_lega = st.session_state.get('ultima_lega_STATUS_WIDGET', 'no tiene ultima lega')
            for p in find_paragraphs_to_remove(doc, s_u_lega, "ULTIMA_LEGA_STATUS"): remove_paragraph(p)


      
            s_renovacion_caa = st.session_state.get('renovacion_caa_STATUS', 'No aplica')
            for p in find_paragraphs_to_remove(doc, s_renovacion_caa, "RENOVACION_CAA_STATUS"): remove_paragraph(p)

            s_lega = st.session_state.get('lega_STATUS', 'Seleccione...')
            for p in find_paragraphs_to_remove(doc, s_lega, "LEGA_STATUS"): remove_paragraph(p)

            s_rree_status = st.session_state.get('RESIDUOS_ESPECIALES_STATUS', 'Seleccione...')
            for p in find_paragraphs_to_remove(doc, s_rree_status, "RESIDUOS_ESPECIALES_STATUS"): remove_paragraph(p)
            
            s_che_status = st.session_state.get('CHE_STATUS', 'Seleccione...')
            for p in find_paragraphs_to_remove(doc, s_che_status, "CHE_STATUS"): remove_paragraph(p)
            
            s_gestion_res = st.session_state.get('GESTION_RESIDUOS_STATUS', 'Correcta')
            for p in find_paragraphs_to_remove(doc, s_gestion_res, "RESIDUOS_GESTION"): remove_paragraph(p)
            
            s_girsu = st.session_state.get('GIRSU_STATUS', 'Aplica')
            for p in find_paragraphs_to_remove(doc, s_girsu, "GIRSU_STATUS"): remove_paragraph(p)
            
            s_patogenicos = st.session_state.get('PATOGENICOS_STATUS', 'No aplica')
            for p in find_paragraphs_to_remove(doc, s_patogenicos, "PATOGENICOS_STATUS"): remove_paragraph(p)

            s_asp = st.session_state.get('ASP_STATUS', 'Seleccione...')
            for p in find_paragraphs_to_remove(doc, s_asp, "ASP_STATUS"): remove_paragraph(p)
            
            s_valvulas = st.session_state.get('VALVULAS_STATUS', 'Cumple')
            for p in find_paragraphs_to_remove(doc, s_valvulas, "VALVULAS_CALIBRACION_STATUS"): remove_paragraph(p)

            s_ada = st.session_state.get('ADA_STATUS', 'Seleccione...')
            for p in find_paragraphs_to_remove(doc, s_ada, "ADA_STATUS"): remove_paragraph(p)
            
            s_renpre = st.session_state.get('RENPRE_STATUS', 'Seleccione...')
            for p in find_paragraphs_to_remove(doc, s_renpre, "RENPRE_STATUS"): remove_paragraph(p)
            
            s_seguro = st.session_state.get('SEGURO_STATUS', 'Vigente')
            for p in find_paragraphs_to_remove(doc, s_seguro, "SEGURO_STATUS"): remove_paragraph(p)

            s_acumar = st.session_state.get('ACUMAR_STATUS', 'Seleccione...')
            for p in find_paragraphs_to_remove(doc, s_acumar, "ACUMAR_STATUS"): remove_paragraph(p)

            s_se = st.session_state.get('SE_STATUS', 'Seleccione...')
            for p in find_paragraphs_to_remove(doc, s_se, "SE_STATUS"): remove_paragraph(p)
            
            s_1102 = st.session_state.get('INSCRIPCION_1102', 'No aplica')
            for p in find_paragraphs_to_remove(doc, s_1102, "INSCRIPCION_1102"): remove_paragraph(p)
            
            s_404 = st.session_state.get('AUDITORIA_404', 'No aplica')
            for p in find_paragraphs_to_remove(doc, s_404, "AUDITORIA_404"): remove_paragraph(p)
            
            s_277 = st.session_state.get('INSCRIPCION_277', 'No aplica')
            for p in find_paragraphs_to_remove(doc, s_277, "INSCRIPCION_277"): remove_paragraph(p)
            status_caa = st.session_state.get('ULTIMO_CAA_STATUS', 'no tiene ultimo caa')
            for p in find_paragraphs_to_remove(doc, status_caa, "ULTIMO_CAA_STATUS"):remove_paragraph(p)

            s_hidraulica = st.session_state.get('HIDRAULICA_STATUS', 'Vigente')
            for p in find_paragraphs_to_remove(doc, s_hidraulica, "HIDRAULICA_STATUS"):remove_paragraph(p)

            s_vuelco = st.session_state.get('VUELCO_STATUS', 'Vigente')
            for p in find_paragraphs_to_remove(doc, s_vuelco, "VUELCO_STATUS"):remove_paragraph(p)

            s_explotacion = st.session_state.get('EXPLOTACION_STATUS', 'Vigente')
            for p in find_paragraphs_to_remove(doc, s_explotacion, "EXPLOTACION_STATUS"): remove_paragraph(p)
            
            s_plan_acumar = st.session_state.get('PLANDEADECUACIONACUMAR_STATUS', 'No aplica Plan de adecuacion')
            for p in find_paragraphs_to_remove(doc, s_plan_acumar, "PLANDEADECUACIONACUMAR_STATUS"): remove_paragraph(p)

           
            s_aysa = st.session_state.get('AYSA_STATUS', 'No aplica')
            for p in find_paragraphs_to_remove(doc, s_aysa, "AYSA_STATUS"): remove_paragraph(p)

            
            s_u_pdeg = st.session_state.get('ULTIMA_LEGA_STATUS_WIDGET', 'Tiene ultima LEGA/PDEG')
            for p in find_paragraphs_to_remove(doc, s_u_pdeg, "ULTIMA_LEGA_STATUS"): 
                remove_paragraph(p)
           
            s_u_hidraulica = st.session_state.get('ULTIMOHIDRAULICA_STATUS', 'tiene ultimo permiso hidraulica')
            for p in find_paragraphs_to_remove(doc, s_u_hidraulica, "ULTIMOHIDRAULICA_STATUS"):
                remove_paragraph(p)
          
            s_u_vuelco = st.session_state.get('ULTIMOVUELCO_STATUS_WIDGET', 'tiene ultimo permiso vuelco')
            for p in find_paragraphs_to_remove(doc, s_u_vuelco, "ULTIMOVUELCO_STATUS"):
                remove_paragraph(p)

            s_u_explotacion = st.session_state.get('ULTIMOEXPLOTACION_STATUS_WIDGET', 'tiene ultimo permiso explotacion')
            for p in find_paragraphs_to_remove(doc, s_u_explotacion, "ULTIMOEXPLOTACION_STATUS"):
                remove_paragraph(p)

            # --- 4. Insert Hallazgos ---
            for i, h_data in enumerate(st.session_state['hallazgos_widgets_list']):
                agregar_hallazgo_formateado_al_doc(
                    doc=doc, 
                    index=i + 1,
                    situacion=h_data.get('situacion', ''),
                    autoridad=h_data.get('autoridad', ''),
                    riesgo=h_data.get('riesgo', ''),
                    recomendacion=h_data.get('recomendacion', ''))

            # 5. Insert Monitoring Table
            insertar_tabla_manual_dinamica(doc, st.session_state['muestreo_filas_datos'])

            # 6. Save to BytesIO object for download
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0) # Rewind the stream to the beginning

            file_name = f"Informe_{user_data['RAZON_SOCIAL']}.docx"
            st.download_button(
                label="Descargar Informe DOCX",
                data=output_stream,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.success("✅ Informe generado y listo para descargar. RECORDA REVISARLO Y AGREGAR LA CONCLUSIONES!!")

        except Exception as e:
            st.error(f"❌ Error crítico durante la generación del informe: {e}")
            st.exception(e)
            st.warning("Por favor, revisa los campos y la plantilla DOCX.")

# Inform the user how to run the Streamlit app
st.success("GENERACION DE INFORMES AMBIENTALES, FEBRERO 2026")
